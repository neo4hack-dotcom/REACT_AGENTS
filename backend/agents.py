from __future__ import annotations

import asyncio
import json
import os
import re
from datetime import datetime
from pathlib import Path
from difflib import SequenceMatcher
from typing import Any, Callable, Dict, List, TypedDict
from urllib.parse import urljoin, urlparse

import feedparser
import requests
from bs4 import BeautifulSoup
from elasticsearch import Elasticsearch
from langchain_core.messages import HumanMessage, SystemMessage
from langgraph.graph import END, StateGraph


class AgentResult(TypedDict, total=False):
    sql: str
    rows: List[Any]
    answer: str
    details: Any


StepCallback = Callable[[Dict[str, Any]], None]
PROJECT_ROOT = Path(__file__).resolve().parent.parent


def _emit_step(on_step: StepCallback | None, step: Dict[str, Any]) -> None:
    if not on_step:
        return
    try:
        on_step(step)
    except Exception:
        # Streaming callbacks should never break agent execution.
        return


def _coerce_message_content(content: Any) -> str:
    if isinstance(content, str):
        return content

    if isinstance(content, list):
        parts: List[str] = []
        for item in content:
            if isinstance(item, str):
                parts.append(item)
            elif isinstance(item, dict):
                if "text" in item:
                    parts.append(str(item["text"]))
                elif "content" in item:
                    parts.append(str(item["content"]))
        return "\n".join(parts)

    return str(content)


def _extract_json_block(text: str) -> str:
    match = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", text)
    return match.group(1) if match else text


def _safe_int(value: Any) -> int | None:
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _normalize_identifier(value: Any) -> str:
    return re.sub(r"[^a-z0-9]+", "", str(value or "").lower())


def _normalize_clickhouse_specific_param_name(value: Any) -> str:
    token = str(value or "").strip().strip("{}")
    if not token:
        return ""
    return re.sub(r"[^A-Za-z0-9_]+", "", token)


def _extract_clickhouse_specific_tokens_from_sql(sql: str) -> List[str]:
    tokens: List[str] = []
    if not sql:
        return tokens

    for match in re.finditer(r"\{([A-Za-z][A-Za-z0-9_]*)\}", sql):
        token = _normalize_clickhouse_specific_param_name(match.group(1))
        if token and token not in tokens:
            tokens.append(token)

    for match in re.finditer(r"(?<![A-Za-z0-9_])(P\d{1,3})(?![A-Za-z0-9_])", sql):
        token = _normalize_clickhouse_specific_param_name(match.group(1))
        if token and token not in tokens:
            tokens.append(token)

    return tokens


def _normalize_clickhouse_specific_templates(config: Dict[str, Any]) -> List[Dict[str, Any]]:
    raw_templates = config.get("query_templates")
    if raw_templates is None:
        raw_templates = config.get("queries")
    if raw_templates is None:
        raw_templates = config.get("templates")
    if raw_templates is None and isinstance(config.get("sql_query_template"), str):
        raw_templates = [
            {
                "name": "default",
                "sql": config.get("sql_query_template"),
                "parameters": config.get("sql_parameters") or [],
            }
        ]

    entries: List[Dict[str, Any]] = []
    if isinstance(raw_templates, dict):
        for name, value in raw_templates.items():
            if isinstance(value, str):
                entries.append({"name": str(name), "sql": value})
            elif isinstance(value, dict):
                entries.append({"name": str(value.get("name") or name), **value})
    elif isinstance(raw_templates, list):
        for idx, item in enumerate(raw_templates):
            if isinstance(item, str):
                entries.append({"name": f"query_{idx + 1}", "sql": item})
            elif isinstance(item, dict):
                entries.append(item)

    normalized: List[Dict[str, Any]] = []
    for idx, item in enumerate(entries):
        sql = str(item.get("sql") or item.get("query") or item.get("template") or "").strip()
        if not sql:
            continue

        name = str(item.get("name") or f"query_{idx + 1}").strip()
        if not name:
            name = f"query_{idx + 1}"
        description = str(item.get("description") or item.get("objective") or "").strip()

        raw_params = item.get("parameters")
        if raw_params is None:
            raw_params = item.get("params")

        parameters: List[Dict[str, Any]] = []
        if isinstance(raw_params, list):
            for raw_param in raw_params:
                if isinstance(raw_param, str):
                    param_name = _normalize_clickhouse_specific_param_name(raw_param)
                    if not param_name:
                        continue
                    parameters.append(
                        {
                            "name": param_name,
                            "required": True,
                            "quote": "auto",
                            "default": None,
                            "description": "",
                        }
                    )
                elif isinstance(raw_param, dict):
                    param_name = _normalize_clickhouse_specific_param_name(raw_param.get("name"))
                    if not param_name:
                        continue
                    quote_mode = str(raw_param.get("quote") or "auto").strip().lower()
                    if quote_mode not in {"auto", "none", "string"}:
                        quote_mode = "auto"
                    parameters.append(
                        {
                            "name": param_name,
                            "required": bool(raw_param.get("required", True)),
                            "quote": quote_mode,
                            "default": raw_param.get("default"),
                            "description": str(raw_param.get("description") or "").strip(),
                        }
                    )

        if len(parameters) == 0:
            derived_tokens = _extract_clickhouse_specific_tokens_from_sql(sql)
            parameters = [
                {
                    "name": token,
                    "required": True,
                    "quote": "auto",
                    "default": None,
                    "description": "",
                }
                for token in derived_tokens
            ]

        normalized.append(
            {
                "name": name,
                "description": description,
                "sql": sql,
                "parameters": parameters,
            }
        )

    return normalized


def _agent_catalog_entry(agent: Dict[str, Any]) -> Dict[str, Any]:
    entry = {
        "id": agent.get("id"),
        "name": agent.get("name"),
        "agent_type": agent.get("agent_type"),
        "role": agent.get("role"),
        "objectives": agent.get("objectives"),
    }

    if str(agent.get("agent_type") or "").strip().lower() == "clickhouse_specific":
        config = build_agent_execution_config(agent)
        templates = _normalize_clickhouse_specific_templates(config)
        entry["query_templates"] = [
            {
                "name": template.get("name"),
                "description": template.get("description"),
                "parameters": [param.get("name") for param in template.get("parameters", []) if param.get("name")],
            }
            for template in templates
        ]
        entry["default_query"] = config.get("default_query")
        entry["supports_params"] = True

    return entry


def _agent_aliases(agent: Dict[str, Any]) -> List[str]:
    aliases: List[str] = []
    name = str(agent.get("name") or "").strip()
    agent_type = str(agent.get("agent_type") or "").strip()
    role = str(agent.get("role") or "").strip()

    if name:
        aliases.append(name)
    if agent_type:
        aliases.append(agent_type)
        aliases.append(agent_type.replace("_", " "))
    if role:
        aliases.append(role)

    # Keep insertion order while removing duplicates.
    deduped: List[str] = []
    seen: set[str] = set()
    for alias in aliases:
        key = alias.lower()
        if key in seen:
            continue
        seen.add(key)
        deduped.append(alias)
    return deduped


def _best_fuzzy_agent_match(query: str, candidates: List[Dict[str, Any]]) -> Dict[str, Any] | None:
    norm_query = _normalize_identifier(query)
    if len(norm_query) < 4:
        return None

    scored: List[tuple[float, Dict[str, Any]]] = []
    for agent in candidates:
        best_for_agent = 0.0
        for alias in _agent_aliases(agent):
            norm_alias = _normalize_identifier(alias)
            if not norm_alias:
                continue
            score = SequenceMatcher(None, norm_query, norm_alias).ratio()
            if score > best_for_agent:
                best_for_agent = score
        if best_for_agent > 0:
            scored.append((best_for_agent, agent))

    if not scored:
        return None

    scored.sort(key=lambda item: item[0], reverse=True)
    best_score, best_agent = scored[0]
    second_score = scored[1][0] if len(scored) > 1 else 0.0

    # Require a strong match. If two candidates are too close, avoid random routing.
    if best_score < 0.84:
        return None
    if (best_score - second_score) < 0.06 and best_score < 0.96:
        return None
    return best_agent


def _resolve_agent_from_call(call: Dict[str, Any], candidates: List[Dict[str, Any]]) -> Dict[str, Any] | None:
    if not candidates:
        return None

    requested_id = _safe_int(call.get("agent_id"))
    if requested_id is not None:
        direct = next((agent for agent in candidates if _safe_int(agent.get("id")) == requested_id), None)
        if direct:
            return direct

    requested_name = str(call.get("agent_name") or call.get("agent") or "").strip()
    requested_type = str(call.get("agent_type") or "").strip()

    if requested_name:
        normalized_name = _normalize_identifier(requested_name)

        # First, strong deterministic matches on aliases (name/type/role).
        for agent in candidates:
            aliases = _agent_aliases(agent)
            for alias in aliases:
                if alias.lower() == requested_name.lower():
                    return agent
                if normalized_name and _normalize_identifier(alias) == normalized_name:
                    return agent

        # Then, relaxed containment on normalized aliases.
        if normalized_name:
            for agent in candidates:
                for alias in _agent_aliases(agent):
                    normalized_alias = _normalize_identifier(alias)
                    if not normalized_alias:
                        continue
                    if normalized_name in normalized_alias or normalized_alias in normalized_name:
                        return agent

        fuzzy_name_match = _best_fuzzy_agent_match(requested_name, candidates)
        if fuzzy_name_match:
            return fuzzy_name_match

    if requested_type:
        normalized_type = _normalize_identifier(requested_type)
        for agent in candidates:
            agent_type = str(agent.get("agent_type") or "").strip()
            if agent_type.lower() == requested_type.lower():
                return agent
            if normalized_type and _normalize_identifier(agent_type) == normalized_type:
                return agent

        fuzzy_type_match = _best_fuzzy_agent_match(requested_type, candidates)
        if fuzzy_type_match:
            return fuzzy_type_match

    if len(candidates) == 1:
        return candidates[0]

    return None


def _extract_urls_from_text(text: str) -> List[str]:
    matches = re.findall(r"https?://[^\s<>\"]+", text or "")
    urls: List[str] = []
    seen: set[str] = set()
    for raw in matches:
        cleaned = str(raw).strip().strip("<>[](){}\"'")
        cleaned = cleaned.rstrip(").,;!?]")
        parsed = urlparse(cleaned)
        if parsed.scheme not in {"http", "https"} or not parsed.netloc:
            continue
        normalized = cleaned.rstrip("/")
        if normalized in seen:
            continue
        seen.add(normalized)
        urls.append(cleaned)
    return urls


def _normalize_navigation_url_candidates(urls: List[str]) -> tuple[List[str], List[str]]:
    valid: List[str] = []
    invalid: List[str] = []
    seen_valid: set[str] = set()
    seen_invalid: set[str] = set()

    for raw in urls:
        candidate = str(raw or "").strip().strip("<>[](){}\"'")
        candidate = candidate.rstrip(").,;!?]")
        if not candidate:
            continue

        if "://" not in candidate:
            bare_candidate = candidate.lstrip(".,;:")
            if re.match(r"^(?:www\.)?[A-Za-z0-9-]+(?:\.[A-Za-z0-9-]+)+(?:[/?#].*)?$", bare_candidate):
                candidate = f"https://{bare_candidate}"

        parsed = urlparse(candidate)
        if parsed.scheme in {"http", "https"} and parsed.netloc:
            normalized = candidate.rstrip("/")
            if normalized in seen_valid:
                continue
            seen_valid.add(normalized)
            valid.append(candidate)
            continue

        normalized_invalid = candidate.lower()
        if normalized_invalid in seen_invalid:
            continue
        seen_invalid.add(normalized_invalid)
        invalid.append(candidate)

    return valid, invalid


def _extract_raw_navigation_tokens(text: str) -> List[str]:
    matches = re.findall(r"https?://[^\s<>\"]*", text or "")
    matches.extend(
        re.findall(
            r"\b(?:www\.)?[A-Za-z0-9-]+(?:\.[A-Za-z0-9-]+)+(?:/[^\s<>\"]*)?\b",
            text or "",
        )
    )
    tokens: List[str] = []
    seen: set[str] = set()
    for match in matches:
        token = str(match or "").strip()
        if not token:
            continue
        key = token.lower()
        if key in seen:
            continue
        seen.add(key)
        tokens.append(token)
    return tokens


def _build_recent_manager_trace(steps: List[Dict[str, Any]], limit: int = 8) -> str:
    recent = steps[-limit:] if len(steps) > limit else steps
    lines: List[str] = []
    for step in recent:
        status = str(step.get("status") or "unknown")
        agent = str(step.get("agent") or "")
        message = str(step.get("message") or "")
        rationale = str(step.get("rationale") or "")
        result = step.get("result")
        result_answer = ""
        if isinstance(result, dict):
            result_answer = str(result.get("answer") or "")
        parts = [f"status={status}"]
        if agent:
            parts.append(f"agent={agent}")
        if message:
            parts.append(f"message={message[:240]}")
        if rationale:
            parts.append(f"rationale={rationale[:240]}")
        if result_answer:
            parts.append(f"result={result_answer[:280]}")
        lines.append("- " + " | ".join(parts))
    return "\n".join(lines) if lines else "- (no prior steps)"


def _plan_revision_summary(previous_plan: str, current_plan: str) -> str:
    previous = str(previous_plan or "").strip()
    current = str(current_plan or "").strip()
    if not previous and not current:
        return "No explicit plan was produced."
    if not previous and current:
        return "Initial plan established."
    if previous == current:
        return "Plan unchanged from previous iteration."
    return "Plan updated based on new evidence."


def _clip_text(value: str, max_chars: int = 1200) -> str:
    text = str(value or "").strip()
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "..."


def _extract_requested_item_limit(text: str, default_limit: int = 10, hard_max: int = 50) -> int:
    lowered = (text or "").lower()
    patterns = [
        r"(?:latest|top|first)\s+(\d{1,3})",
        r"(\d{1,3})\s+(?:articles|results|links|titles|items)",
    ]
    for pattern in patterns:
        match = re.search(pattern, lowered)
        if not match:
            continue
        try:
            value = int(match.group(1))
            return max(1, min(value, hard_max))
        except Exception:
            continue
    return max(1, min(default_limit, hard_max))


def _is_probable_article_anchor(title: str, url: str) -> bool:
    title_norm = " ".join(title.split()).strip()
    if len(title_norm) < 8:
        return False

    lower_title = title_norm.lower()
    if lower_title in {"menu", "accueil", "home", "login", "connexion"}:
        return False

    lower_url = url.lower()
    blocked_tokens = [
        "/login",
        "/connexion",
        "/account",
        "/compte",
        "/abonnement",
        "/subscribe",
        "/newsletter",
        "/video",
        "/podcast",
        "/tag/",
        "#",
        "javascript:",
        "mailto:",
    ]
    if any(token in lower_url for token in blocked_tokens):
        return False

    return True


def _extract_article_links_from_html(
    page_url: str,
    html: str,
    max_links: int,
    same_domain_only: bool,
) -> List[Dict[str, str]]:
    soup = BeautifulSoup(html, "html.parser")
    base_domain = urlparse(page_url).netloc.lower()

    selectors = [
        "article a[href]",
        "main a[href]",
        "h1 a[href]",
        "h2 a[href]",
        "h3 a[href]",
        "[data-testid] a[href]",
        "a[href]",
    ]

    links: List[Dict[str, str]] = []
    seen_urls: set[str] = set()

    for selector in selectors:
        for anchor in soup.select(selector):
            href = str(anchor.get("href") or "").strip()
            if not href:
                continue
            if href.startswith("#") or href.lower().startswith("javascript:") or href.lower().startswith("mailto:"):
                continue

            absolute_url = urljoin(page_url, href)
            parsed = urlparse(absolute_url)
            if parsed.scheme not in {"http", "https"}:
                continue

            if same_domain_only and base_domain:
                target_domain = parsed.netloc.lower()
                if target_domain and target_domain != base_domain and not target_domain.endswith(f".{base_domain}"):
                    continue

            title = anchor.get_text(" ", strip=True)
            if not title:
                continue
            title = " ".join(title.split())
            if not _is_probable_article_anchor(title, absolute_url):
                continue

            normalized_url = absolute_url.rstrip("/")
            if normalized_url in seen_urls:
                continue
            seen_urls.add(normalized_url)
            links.append({"title": title, "url": absolute_url})

            if len(links) >= max_links:
                return links

    return links


def _discover_feed_links_from_html(page_url: str, html: str, max_links: int = 8) -> List[str]:
    soup = BeautifulSoup(html or "", "html.parser")
    discovered: List[str] = []
    seen: set[str] = set()

    selectors = [
        "link[rel='alternate'][type*='rss']",
        "link[rel='alternate'][type*='atom']",
        "link[rel='alternate'][type*='xml']",
        "a[href*='rss']",
        "a[href*='feed']",
    ]

    for selector in selectors:
        for node in soup.select(selector):
            href = str(node.get("href") or "").strip()
            if not href:
                continue
            candidate = urljoin(page_url, href).strip()
            parsed = urlparse(candidate)
            if parsed.scheme not in {"http", "https"} or not parsed.netloc:
                continue
            normalized = candidate.rstrip("/")
            if normalized in seen:
                continue
            seen.add(normalized)
            discovered.append(candidate)
            if len(discovered) >= max_links:
                return discovered

    return discovered


def _build_common_feed_candidates(url: str, max_links: int = 8) -> List[str]:
    parsed = urlparse(str(url or "").strip())
    if parsed.scheme not in {"http", "https"} or not parsed.netloc:
        return []
    root = f"{parsed.scheme}://{parsed.netloc}"
    candidates = [
        f"{root}/rss",
        f"{root}/rss.xml",
        f"{root}/feed",
        f"{root}/feed.xml",
        f"{root}/atom.xml",
        f"{root}/feeds/all.atom.xml",
        f"{root}/actualites/rss.xml",
        f"{root}/news/rss.xml",
    ]
    deduped: List[str] = []
    seen: set[str] = set()
    for candidate in candidates:
        normalized = candidate.rstrip("/")
        if normalized in seen:
            continue
        seen.add(normalized)
        deduped.append(candidate)
        if len(deduped) >= max_links:
            break
    return deduped


def _build_www_url_variants(url: str) -> List[str]:
    raw = str(url or "").strip()
    parsed = urlparse(raw)
    if parsed.scheme not in {"http", "https"} or not parsed.netloc:
        return []

    variants: List[str] = []
    host = parsed.netloc
    path = parsed.path or "/"
    suffix = f"{path}{'?' + parsed.query if parsed.query else ''}{'#' + parsed.fragment if parsed.fragment else ''}"

    if host.startswith("www."):
        no_www = host[4:]
        variants.append(f"{parsed.scheme}://{no_www}{suffix}")
    else:
        variants.append(f"{parsed.scheme}://www.{host}{suffix}")

    deduped: List[str] = []
    seen: set[str] = set()
    for item in variants:
        normalized = item.rstrip("/")
        if normalized in seen:
            continue
        seen.add(normalized)
        deduped.append(item)
    return deduped


def _build_google_news_domain_feed_candidates(url: str, max_links: int = 2) -> List[str]:
    parsed = urlparse(str(url or "").strip())
    if parsed.scheme not in {"http", "https"} or not parsed.netloc:
        return []
    domain = parsed.netloc.lower()
    if domain.startswith("www."):
        domain = domain[4:]
    if not domain:
        return []

    queries = [
        f"site:{domain}",
        domain,
    ]
    candidates: List[str] = []
    for query in queries:
        candidates.append(
            f"https://news.google.com/rss/search?q={query}&hl=fr&gl=FR&ceid=FR:fr"
        )
        candidates.append(
            f"https://news.google.com/rss/search?q={query}&hl=en-US&gl=US&ceid=US:en"
        )

    deduped: List[str] = []
    seen: set[str] = set()
    for candidate in candidates:
        normalized = candidate.rstrip("/")
        if normalized in seen:
            continue
        seen.add(normalized)
        deduped.append(candidate)
        if len(deduped) >= max_links:
            break
    return deduped


def _build_reader_proxy_candidates(url: str, max_links: int = 2) -> List[str]:
    raw = str(url or "").strip()
    if not raw:
        return []
    if raw.startswith("https://r.jina.ai/"):
        return []
    parsed = urlparse(raw)
    if parsed.scheme not in {"http", "https"} or not parsed.netloc:
        return []

    normalized = f"{parsed.netloc}{parsed.path or '/'}"
    if parsed.query:
        normalized += f"?{parsed.query}"
    candidates = [
        f"https://r.jina.ai/http://{normalized}",
        f"https://r.jina.ai/http://{parsed.netloc}",
    ]
    deduped: List[str] = []
    seen: set[str] = set()
    for candidate in candidates:
        if candidate in seen:
            continue
        seen.add(candidate)
        deduped.append(candidate)
        if len(deduped) >= max_links:
            break
    return deduped


def _build_http_headers(
    *,
    language_hint: str = "en-US,en;q=0.9,fr;q=0.8",
    prefer_xml: bool = False,
    referer: str | None = None,
    user_agent: str | None = None,
) -> Dict[str, str]:
    headers: Dict[str, str] = {
        "User-Agent": user_agent
        or (
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/123.0.0.0 Safari/537.36"
        ),
        "Accept-Language": language_hint,
        "Cache-Control": "no-cache",
        "Pragma": "no-cache",
    }
    if prefer_xml:
        headers["Accept"] = (
            "application/rss+xml, application/atom+xml, application/xml, text/xml;q=0.9, */*;q=0.5"
        )
    else:
        headers["Accept"] = (
            "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8"
        )
    if referer:
        headers["Referer"] = referer
    return headers


def _resolve_accept_language(language_hint: Any) -> str:
    raw = str(language_hint or "").strip().lower()
    if not raw:
        return "en-US,en;q=0.9,fr;q=0.8"
    if "," in raw or ";" in raw:
        return str(language_hint)
    if raw in {"fr", "fr-fr"}:
        return "fr-FR,fr;q=0.95,en;q=0.7"
    if raw in {"en", "en-us", "en-gb"}:
        return "en-US,en;q=0.95,fr;q=0.6"
    return f"{raw};q=1.0,en;q=0.7"


def _fetch_url_with_fallback(
    url: str,
    *,
    timeout_seconds: int = 20,
    prefer_xml: bool = False,
    language_hint: str = "en-US,en;q=0.9,fr;q=0.8",
    allow_common_feed_candidates: bool = False,
    allow_google_news_fallback: bool = False,
    allow_reader_proxy_fallback: bool = True,
) -> tuple[requests.Response | None, List[str]]:
    raw_url = str(url or "").strip()
    parsed = urlparse(raw_url)
    if parsed.scheme not in {"http", "https"} or not parsed.netloc:
        return None, [f"{raw_url} -> invalid_url"]

    primary_candidates: List[str] = [raw_url]
    primary_candidates.extend(_build_www_url_variants(raw_url))
    if allow_common_feed_candidates:
        primary_candidates.extend(_build_common_feed_candidates(raw_url, max_links=8))
    if allow_google_news_fallback:
        primary_candidates.extend(_build_google_news_domain_feed_candidates(raw_url, max_links=4))

    proxy_candidates: List[str] = []
    if allow_reader_proxy_fallback:
        for candidate in primary_candidates[:6]:
            proxy_candidates.extend(_build_reader_proxy_candidates(candidate, max_links=2))

    candidates: List[str] = []
    seen_candidates: set[str] = set()
    for candidate in [*primary_candidates, *proxy_candidates]:
        normalized = str(candidate or "").strip()
        if not normalized:
            continue
        key = normalized.rstrip("/")
        if key in seen_candidates:
            continue
        seen_candidates.add(key)
        candidates.append(normalized)

    tried: set[str] = set()
    failures: List[str] = []
    user_agents = [
        (
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/123.0.0.0 Safari/537.36"
        ),
        (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:123.0) "
            "Gecko/20100101 Firefox/123.0"
        ),
    ]

    for candidate in candidates:
        normalized_candidate = str(candidate or "").strip()
        if not normalized_candidate:
            continue
        if normalized_candidate in tried:
            continue
        tried.add(normalized_candidate)
        candidate_parsed = urlparse(normalized_candidate)
        candidate_referer = ""
        if candidate_parsed.scheme in {"http", "https"} and candidate_parsed.netloc:
            candidate_referer = f"{candidate_parsed.scheme}://{candidate_parsed.netloc}/"
        elif parsed.scheme in {"http", "https"} and parsed.netloc:
            candidate_referer = f"{parsed.scheme}://{parsed.netloc}/"

        for ua in user_agents:
            headers = _build_http_headers(
                language_hint=_resolve_accept_language(language_hint),
                prefer_xml=prefer_xml,
                referer=candidate_referer,
                user_agent=ua,
            )
            try:
                response = requests.get(
                    normalized_candidate,
                    headers=headers,
                    timeout=timeout_seconds,
                    allow_redirects=True,
                )
                if response.ok and response.content:
                    return response, failures
                failures.append(f"{normalized_candidate} -> HTTP {response.status_code}")
            except Exception as exc:
                failures.append(f"{normalized_candidate} -> {exc}")

    return None, failures


def _extract_feed_entries_from_content(content: bytes | str | None) -> List[Dict[str, Any]]:
    if not content:
        return []
    try:
        parsed = feedparser.parse(content)
        entries = parsed.entries or []
        normalized: List[Dict[str, Any]] = []
        for entry in entries:
            if isinstance(entry, dict):
                normalized.append(entry)
            else:
                try:
                    normalized.append(dict(entry))
                except Exception:
                    continue
        return normalized
    except Exception:
        return []


def _entry_to_link_item(entry: Dict[str, Any], base_url: str) -> Dict[str, str] | None:
    title = " ".join(str(entry.get("title") or "").split()).strip()
    link = str(entry.get("link") or entry.get("id") or "").strip()
    if not title or not link:
        return None
    resolved = urljoin(base_url, link)
    parsed = urlparse(resolved)
    if parsed.scheme not in {"http", "https"} or not parsed.netloc:
        return None
    return {"title": title, "url": resolved}


def _extract_fallback_links_from_html(
    page_url: str,
    html: str,
    max_links: int,
    same_domain_only: bool,
) -> List[Dict[str, str]]:
    soup = BeautifulSoup(html or "", "html.parser")
    base_domain = urlparse(page_url).netloc.lower()
    links: List[Dict[str, str]] = []
    seen_urls: set[str] = set()

    for anchor in soup.select("a[href]"):
        href = str(anchor.get("href") or "").strip()
        if not href:
            continue
        if href.startswith("#") or href.lower().startswith("javascript:") or href.lower().startswith("mailto:"):
            continue

        absolute_url = urljoin(page_url, href)
        parsed = urlparse(absolute_url)
        if parsed.scheme not in {"http", "https"} or not parsed.netloc:
            continue

        if same_domain_only and base_domain:
            target_domain = parsed.netloc.lower()
            if target_domain and target_domain != base_domain and not target_domain.endswith(f".{base_domain}"):
                continue

        title = anchor.get_text(" ", strip=True)
        if not title:
            title = str(anchor.get("title") or parsed.path or parsed.netloc or "").strip()
        title = " ".join(title.split())
        if len(title) < 3:
            continue

        normalized_url = absolute_url.rstrip("/")
        if normalized_url in seen_urls:
            continue
        seen_urls.add(normalized_url)
        links.append({"title": title, "url": absolute_url})
        if len(links) >= max_links:
            break

    return links


def _extract_links_from_text_blob(
    page_url: str,
    text: str,
    max_links: int,
    same_domain_only: bool,
) -> List[Dict[str, str]]:
    base_domain = urlparse(page_url).netloc.lower()
    matches = re.findall(r"https?://[^\s<>\"]+", text or "")
    links: List[Dict[str, str]] = []
    seen: set[str] = set()
    for raw in matches:
        candidate = str(raw or "").strip().strip("<>[](){}\"'")
        candidate = candidate.rstrip(").,;!?]")
        parsed = urlparse(candidate)
        if parsed.scheme not in {"http", "https"} or not parsed.netloc:
            continue
        if same_domain_only and base_domain:
            target_domain = parsed.netloc.lower()
            if target_domain and target_domain != base_domain and not target_domain.endswith(f".{base_domain}"):
                continue
        normalized = candidate.rstrip("/")
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        title = parsed.path.strip("/") or parsed.netloc
        title = title.replace("-", " ").replace("_", " ")
        title = " ".join(title.split())[:120] or parsed.netloc
        links.append({"title": title, "url": candidate})
        if len(links) >= max_links:
            break
    return links


def _extract_possible_json(text: str) -> Any:
    candidate = _extract_json_block(text or "").strip()
    if not candidate:
        return None
    try:
        return json.loads(candidate)
    except Exception:
        return None


def _is_virtual_data_path(raw_path: str) -> bool:
    normalized = str(raw_path or "").replace("\\", "/").strip()
    return normalized == "/data" or normalized.startswith("/data/")


def _resolve_workspace_storage_path(raw_path: str, fallback_dir: str | None = None) -> str:
    value = os.path.expandvars(os.path.expanduser(str(raw_path or "").strip()))
    if not value:
        base = fallback_dir or str(PROJECT_ROOT)
        return os.path.abspath(base)

    if _is_virtual_data_path(value):
        return os.path.abspath(os.path.join(str(PROJECT_ROOT), value.lstrip("/")))

    if os.path.isabs(value):
        return os.path.abspath(value)

    base = fallback_dir or str(PROJECT_ROOT)
    return os.path.abspath(os.path.join(base, value))


def _resolve_managed_root_path(
    config: Dict[str, Any],
    *,
    default_subdir: str,
) -> str:
    configured_folder = str(config.get("folder_path") or "").strip()
    default_folder = str(PROJECT_ROOT / default_subdir)
    if configured_folder:
        resolved_folder = _resolve_workspace_storage_path(configured_folder, fallback_dir=str(PROJECT_ROOT))
    else:
        resolved_folder = os.path.abspath(default_folder)
    auto_create_folder = bool(config.get("auto_create_folder", True))

    if os.path.exists(resolved_folder):
        if not os.path.isdir(resolved_folder):
            raise ValueError(f"Configured folder is not a directory: {resolved_folder}")
    elif auto_create_folder:
        os.makedirs(resolved_folder, exist_ok=True)
    else:
        raise ValueError(f"Configured folder does not exist: {resolved_folder}")

    return os.path.abspath(resolved_folder)


def _to_relative_path(root_path: str, target_path: str) -> str:
    root = Path(root_path).resolve()
    target = Path(target_path).resolve()
    try:
        relative = target.relative_to(root)
        text = str(relative)
        return text if text else "."
    except Exception:
        return str(target)


def _resolve_safe_managed_file_path(
    *,
    root_path: str,
    raw_path: str,
    required_suffix: str | None = None,
    allowed_suffixes: tuple[str, ...] | None = None,
    allow_outside_folder: bool = False,
) -> str:
    resolved_candidate = _resolve_workspace_storage_path(raw_path, fallback_dir=root_path)
    candidate = Path(resolved_candidate)

    if allowed_suffixes:
        normalized_allowed = tuple(str(item).lower() for item in allowed_suffixes if item)
        if not normalized_allowed:
            normalized_allowed = None
        if normalized_allowed:
            if not candidate.suffix:
                candidate = candidate.with_suffix(normalized_allowed[0])
            elif candidate.suffix.lower() not in normalized_allowed:
                joined = ", ".join(normalized_allowed)
                raise ValueError(f"Expected one of ({joined}) file extensions: {candidate}")
    elif required_suffix:
        normalized_suffix = required_suffix.lower()
        if not candidate.suffix:
            candidate = candidate.with_suffix(normalized_suffix)
        elif candidate.suffix.lower() != normalized_suffix:
            raise ValueError(f"Expected a '{normalized_suffix}' file path: {candidate}")

    resolved = candidate.resolve()
    root = Path(root_path).resolve()
    if not allow_outside_folder:
        try:
            resolved.relative_to(root)
        except Exception as exc:
            raise ValueError(f"Path '{resolved}' is outside the configured folder '{root}'.") from exc

    return str(resolved)


def _is_instructional_excel_workbook_value(value: str) -> bool:
    raw = str(value or "").strip()
    if not raw:
        return False
    lowered = raw.lower()
    if re.search(r"\.(xlsx|xlsm|xltx|xltm)$", lowered):
        return False
    if "/" in raw or "\\" in raw:
        return False
    word_count = len(re.findall(r"[A-Za-zÀ-ÿ0-9_'-]+", raw))
    instruction_markers = (
        "ia",
        "ai",
        "decide",
        "décide",
        "choisit",
        "choose",
        "context",
        "contexte",
        "folder",
        "dossier",
    )
    return word_count >= 5 and (len(raw) > 30 or any(marker in lowered for marker in instruction_markers))


def _normalize_excel_workbook_preference(value: str) -> str:
    raw = str(value or "").strip()
    if not raw:
        return ""
    if _is_instructional_excel_workbook_value(raw):
        return ""
    return raw


def _excel_manager_uses_auto_workbook_selection(config: Dict[str, Any]) -> bool:
    mode = str(config.get("workbook_path_mode") or "").strip().lower()
    if mode in {"auto", "auto_in_folder", "dynamic", "llm_decides", "ia_decide"}:
        return True
    configured = str(config.get("workbook_path") or "").strip()
    return _is_instructional_excel_workbook_value(configured)


def _list_excel_workbooks_in_root(root_path: str) -> List[Path]:
    root = Path(root_path)
    suffixes = {".xlsx", ".xlsm", ".xltx", ".xltm"}
    workbooks: List[Path] = []
    try:
        for item in root.iterdir():
            if not item.is_file():
                continue
            if item.suffix.lower() not in suffixes:
                continue
            workbooks.append(item)
    except Exception:
        return []

    workbooks.sort(key=lambda path: path.stat().st_mtime, reverse=True)
    return workbooks


def _choose_auto_excel_workbook_candidate(
    root_path: str,
    action_name: str,
    allow_overwrite: bool,
) -> str:
    existing = _list_excel_workbooks_in_root(root_path)
    normalized_action = str(action_name or "").strip().lower()

    if normalized_action == "create_workbook":
        if existing and allow_overwrite:
            return str(existing[0])
        if not allow_overwrite:
            index = 1
            while True:
                candidate = Path(root_path) / f"workbook_{index}.xlsx"
                if not candidate.exists():
                    return str(candidate)
                index += 1

    if existing:
        return str(existing[0])

    return str(Path(root_path) / "workbook_auto.xlsx")


def _is_instructional_excel_sheet_value(value: str) -> bool:
    raw = str(value or "").strip()
    if not raw:
        return False
    if len(raw) <= 31 and re.fullmatch(r"[A-Za-z0-9 _-]{1,31}", raw):
        return False
    lowered = raw.lower()
    markers = ("ia", "ai", "decide", "décide", "context", "contexte", "choose", "choisit")
    word_count = len(re.findall(r"[A-Za-zÀ-ÿ0-9_'-]+", raw))
    return word_count >= 4 and any(marker in lowered for marker in markers)


def _normalize_excel_sheet_preference(value: str) -> str:
    raw = str(value or "").strip()
    if not raw:
        return ""
    if _is_instructional_excel_sheet_value(raw):
        return ""
    return raw


def _excel_manager_uses_auto_sheet_selection(config: Dict[str, Any]) -> bool:
    mode = str(config.get("default_sheet_mode") or "").strip().lower()
    if mode in {"auto", "auto_from_context", "dynamic", "llm_decides", "ia_decide"}:
        return True
    configured = str(config.get("default_sheet") or "").strip()
    return _is_instructional_excel_sheet_value(configured)


def _resolve_excel_default_sheet(config: Dict[str, Any]) -> str:
    normalized = _normalize_excel_sheet_preference(str(config.get("default_sheet") or ""))
    return normalized or "Sheet1"


def _resolve_excel_sheet_name_for_action(
    workbook: Any,
    action: Dict[str, Any],
    default_sheet: str,
    auto_sheet_selection: bool,
    action_name: str,
) -> str:
    explicit_sheet = _normalize_excel_sheet_preference(str(action.get("sheet") or ""))
    if explicit_sheet:
        return explicit_sheet[:31] or "Sheet1"

    sheet_names = list(getattr(workbook, "sheetnames", []) or [])
    if auto_sheet_selection and len(sheet_names) > 0:
        active_name = ""
        try:
            active_name = str(workbook.active.title or "").strip()
        except Exception:
            active_name = ""
        if active_name and active_name in sheet_names:
            return active_name[:31]
        return str(sheet_names[0])[:31]

    if str(action_name or "").strip().lower() == "create_sheet" and len(sheet_names) > 0:
        base = (default_sheet or "Sheet1").strip()[:31] or "Sheet1"
        if base not in sheet_names:
            return base
        index = 2
        while True:
            suffix = f"_{index}"
            candidate = f"{base[: max(1, 31 - len(suffix))]}{suffix}"
            if candidate not in sheet_names:
                return candidate
            index += 1

    return (default_sheet or "Sheet1").strip()[:31] or "Sheet1"


def _resolve_excel_workbook_path(
    config: Dict[str, Any],
    workbook_path: str | None = None,
    *,
    action_name: str | None = None,
) -> str:
    configured_workbook = _normalize_excel_workbook_preference(str(config.get("workbook_path") or ""))
    requested_workbook = _normalize_excel_workbook_preference(str(workbook_path or ""))
    root_path = _resolve_managed_root_path(config, default_subdir="data/spreadsheets")

    auto_workbook_selection = _excel_manager_uses_auto_workbook_selection(config)
    if requested_workbook:
        raw_path = requested_workbook
    elif auto_workbook_selection:
        raw_path = _choose_auto_excel_workbook_candidate(
            root_path=root_path,
            action_name=str(action_name or ""),
            allow_overwrite=bool(config.get("allow_overwrite")),
        )
    else:
        raw_path = configured_workbook or "data.xlsx"

    resolved = _resolve_safe_managed_file_path(
        root_path=root_path,
        raw_path=raw_path,
        allowed_suffixes=(".xlsx", ".xlsm", ".xltx", ".xltm"),
        allow_outside_folder=bool(config.get("allow_outside_folder")),
    )
    parent = os.path.dirname(resolved) or root_path
    if bool(config.get("auto_create_folder", True)):
        os.makedirs(parent, exist_ok=True)
    return resolved


def _guess_sheet_name_from_text(text: str, default_name: str) -> str:
    patterns = [
        r"(?:sheet|onglet|feuille)\s+[\"']([^\"']+)[\"']",
        r"(?:sheet|onglet|feuille)\s+([A-Za-z0-9 _-]{1,50})",
    ]
    lowered = text or ""
    for pattern in patterns:
        match = re.search(pattern, lowered, flags=re.IGNORECASE)
        if match:
            value = " ".join(str(match.group(1)).split()).strip()
            if value:
                return value
    return default_name


def _guess_sheet_rename_from_text(text: str) -> tuple[str, str] | None:
    patterns = [
        r"(?:rename|renomme(?:r)?|renommer)\s+(?:the\s+|la\s+|le\s+|l')?(?:sheet|onglet|feuille)\s+[\"']([^\"']+)[\"']\s+(?:to|en|vers)\s+[\"']([^\"']+)[\"']",
        r"(?:rename|renomme(?:r)?|renommer)\s+(?:the\s+|la\s+|le\s+|l')?(?:sheet|onglet|feuille)\s+([A-Za-z0-9 _-]{1,50}?)\s+(?:to|en|vers)\s+([A-Za-z0-9 _-]{1,50})",
    ]
    for pattern in patterns:
        match = re.search(pattern, text or "", flags=re.IGNORECASE)
        if not match:
            continue
        source = " ".join(str(match.group(1)).split()).strip()
        target = " ".join(str(match.group(2)).split()).strip()
        if source and target:
            return source, target
    return None


def _guess_workbook_path_from_text(text: str) -> str | None:
    quoted = re.search(
        r"[\"']([^\"']+\.(?:xlsx|xlsm|xltx|xltm))[\"']",
        text or "",
        flags=re.IGNORECASE,
    )
    if quoted:
        return str(quoted.group(1)).strip()

    bare = re.search(
        r"([A-Za-z]:[\\/][^\s\"']+\.(?:xlsx|xlsm|xltx|xltm)|\.{0,2}[\\/][^\s\"']+\.(?:xlsx|xlsm|xltx|xltm)|[A-Za-z0-9._/-]+\.(?:xlsx|xlsm|xltx|xltm))",
        text or "",
        flags=re.IGNORECASE,
    )
    if bare:
        return str(bare.group(1)).strip()
    return None


def _guess_cell_and_value_from_text(text: str) -> tuple[str, Any] | None:
    patterns = [
        r"(?:cellule|cell)\s+([A-Za-z]{1,3}\d{1,6})\s*(?:=|:|to|avec|with)\s*[\"']([^\"']*)[\"']",
        r"([A-Za-z]{1,3}\d{1,6})\s*(?:=|:)\s*[\"']([^\"']*)[\"']",
        r"(?:cellule|cell)\s+([A-Za-z]{1,3}\d{1,6})\s*(?:=|:|to|avec|with)\s*([^\n]+)",
    ]
    for pattern in patterns:
        match = re.search(pattern, text or "", flags=re.IGNORECASE)
        if not match:
            continue
        cell = str(match.group(1)).upper().strip()
        value = str(match.group(2)).strip()
        if cell:
            return (cell, value)
    return None


def _normalize_excel_action_payload(action: Dict[str, Any]) -> Dict[str, Any]:
    normalized: Dict[str, Any] = dict(action)
    raw_action_name = str(normalized.get("action") or "").strip().lower()
    action_aliases = {
        "create_file": "create_workbook",
        "new_workbook": "create_workbook",
        "remove_workbook": "delete_workbook",
        "delete_file": "delete_workbook",
        "remove_sheet": "delete_sheet",
        "new_sheet": "create_sheet",
        "add_sheet": "create_sheet",
        "write_cell": "write_cells",
        "set_cell": "write_cells",
        "append_row": "append_rows",
        "read_rows": "read_sheet",
    }
    normalized["action"] = action_aliases.get(raw_action_name, raw_action_name)

    if normalized.get("workbook_path") in (None, ""):
        for key in ("file_path", "file", "path", "workbook", "workbook_file"):
            value = normalized.get(key)
            if isinstance(value, str) and value.strip():
                normalized["workbook_path"] = value.strip()
                break

    if normalized.get("sheet") in (None, ""):
        for key in ("sheet_name", "worksheet", "tab", "onglet", "feuille"):
            value = normalized.get(key)
            if isinstance(value, str) and value.strip():
                normalized["sheet"] = value.strip()
                break

    if normalized.get("new_name") in (None, ""):
        for key in ("target_sheet", "new_sheet", "target", "destination", "to"):
            value = normalized.get(key)
            if isinstance(value, str) and value.strip():
                normalized["new_name"] = value.strip()
                break

    if normalized.get("action") == "write_cells" and not isinstance(normalized.get("cells"), dict):
        cell = normalized.get("cell")
        if cell is not None:
            normalized["cells"] = {str(cell): normalized.get("value")}

    if normalized.get("action") == "append_rows" and not isinstance(normalized.get("rows"), list):
        alt_rows = normalized.get("data") or normalized.get("values")
        if isinstance(alt_rows, list):
            normalized["rows"] = alt_rows

    if normalized.get("auto_create_workbook") in (None, ""):
        for key in ("auto_create", "create_if_missing", "create_workbook_if_missing"):
            if key in normalized:
                normalized["auto_create_workbook"] = normalized.get(key)
                break

    return normalized


def _normalize_excel_actions(
    user_input: str,
    config: Dict[str, Any],
    allow_heuristics: bool = True,
) -> List[Dict[str, Any]] | None:
    parsed = _extract_possible_json(user_input)
    if isinstance(parsed, dict):
        if isinstance(parsed.get("actions"), list):
            actions = [_normalize_excel_action_payload(action) for action in parsed["actions"] if isinstance(action, dict)]
            return actions if actions else None
        if isinstance(parsed.get("action"), str):
            return [_normalize_excel_action_payload(parsed)]
    if isinstance(parsed, list):
        actions = [_normalize_excel_action_payload(action) for action in parsed if isinstance(action, dict)]
        if actions:
            return actions

    if not allow_heuristics:
        return None

    text = (user_input or "").strip()
    lowered = text.lower()
    default_sheet = _resolve_excel_default_sheet(config)
    workbook_hint = _guess_workbook_path_from_text(text)
    heuristics: List[Dict[str, Any]] = []

    def add_action(action: Dict[str, Any]) -> None:
        if workbook_hint and not action.get("workbook_path"):
            action["workbook_path"] = workbook_hint
        heuristics.append(_normalize_excel_action_payload(action))

    if re.search(r"(list|show|affiche|liste).*(sheet|sheets|onglet|onglets|feuille|feuilles)", lowered):
        add_action({"action": "list_sheets"})
    elif re.search(r"(create|new|cr[eé]e|créer|ajoute|add).*(workbook|excel|classeur|fichier)", lowered):
        add_action({"action": "create_workbook"})
    elif re.search(r"(delete|remove|supprime|supprimer).*(workbook|excel|classeur|fichier)", lowered):
        add_action({"action": "delete_workbook"})
    elif re.search(r"(rename|renomme|renommer).*(sheet|onglet|feuille)", lowered):
        rename_pair = _guess_sheet_rename_from_text(text)
        if rename_pair:
            source, target = rename_pair
        else:
            source = _guess_sheet_name_from_text(text, default_sheet)
            target_match = re.search(
                r"(?:to|en|vers)\s+[\"']([^\"']+)[\"']|(?:to|en|vers)\s+([A-Za-z0-9 _-]{1,50})",
                text,
                flags=re.IGNORECASE,
            )
            target = source
            if target_match:
                target = (target_match.group(1) or target_match.group(2) or source).strip()
        add_action({"action": "rename_sheet", "sheet": source, "new_name": target})
    elif re.search(r"(create|new|add|cr[eé]e|créer|ajoute).*(sheet|onglet|feuille)", lowered):
        add_action({"action": "create_sheet", "sheet": _guess_sheet_name_from_text(text, default_sheet)})
    elif re.search(r"(delete|remove|supprime|supprimer).*(sheet|onglet|feuille)", lowered):
        add_action({"action": "delete_sheet", "sheet": _guess_sheet_name_from_text(text, default_sheet)})
    elif re.search(r"(read|lire|show|affiche).*(sheet|onglet|feuille)", lowered):
        add_action(
            {
                "action": "read_sheet",
                "sheet": _guess_sheet_name_from_text(text, default_sheet),
                "max_rows": int(config.get("max_rows_read") or 100),
            }
        )
    else:
        cell_value = _guess_cell_and_value_from_text(text)
        if cell_value:
            cell, value = cell_value
            add_action(
                {
                    "action": "write_cells",
                    "sheet": _guess_sheet_name_from_text(text, default_sheet),
                    "cells": {cell: value},
                }
            )

    return heuristics if heuristics else None


async def _execute_excel_actions(
    user_input: str,
    config: Dict[str, Any],
    on_step: StepCallback | None,
    agent_label: str,
    allow_heuristics: bool = True,
) -> AgentResult | None:
    actions = _normalize_excel_actions(user_input, config, allow_heuristics=allow_heuristics)
    if not actions:
        return None

    try:
        import openpyxl  # type: ignore
    except Exception as exc:
        return {
            "answer": (
                "Excel manager requires the `openpyxl` dependency. "
                f"Install it to enable workbook operations. ({exc})"
            )
        }

    default_sheet = _resolve_excel_default_sheet(config)
    auto_create_workbook = bool(config.get("auto_create_workbook"))
    allow_overwrite = bool(config.get("allow_overwrite"))
    max_rows_read = int(config.get("max_rows_read") or 100)
    auto_sheet_selection = _excel_manager_uses_auto_sheet_selection(config)
    lowered_request = (user_input or "").lower()
    request_forces_auto_create = bool(
        re.search(r"auto[_\s-]*create[_\s-]*workbook", lowered_request)
        or re.search(r"(create|new|cr[eé]e|créer).*(workbook|excel|classeur|fichier)", lowered_request)
        or "if absent" in lowered_request
        or "if missing" in lowered_request
    )

    operation_logs: List[Dict[str, Any]] = []
    last_read_rows: List[List[Any]] = []
    last_sheet_names: List[str] = []
    workbook_path_cache: str | None = None

    for action in actions:
        action_name = str(action.get("action") or "").strip().lower()
        action_auto_create_workbook = _coerce_bool(
            action.get("auto_create_workbook"),
            default=(auto_create_workbook or request_forces_auto_create),
        )
        try:
            workbook_path = _resolve_excel_workbook_path(
                config,
                action.get("workbook_path"),
                action_name=action_name,
            )
        except Exception as exc:
            return {
                "answer": f"Unable to resolve workbook path: {exc}",
                "details": {"action": action_name, "requested_workbook_path": action.get("workbook_path")},
            }
        workbook_path_cache = workbook_path
        _emit_step(
            on_step,
            {
                "status": "excel_action_started",
                "agent": agent_label,
                "action": action_name,
                "workbook_path": workbook_path,
            },
        )

        if action_name == "create_workbook":
            if os.path.exists(workbook_path) and not allow_overwrite:
                operation_logs.append(
                    {
                        "action": action_name,
                        "status": "skipped",
                        "reason": "Workbook already exists and overwrite is disabled.",
                        "workbook_path": workbook_path,
                    }
                )
                _emit_step(
                    on_step,
                    {
                        "status": "excel_action_completed",
                        "agent": agent_label,
                        "action": action_name,
                        "result": "skipped_existing",
                    },
                )
                continue

            workbook = openpyxl.Workbook()
            active = workbook.active
            initial_sheet_name = _normalize_excel_sheet_preference(str(action.get("sheet") or "")) or default_sheet
            active.title = initial_sheet_name[:31] or "Sheet1"
            workbook.save(workbook_path)
            operation_logs.append({"action": action_name, "status": "ok", "workbook_path": workbook_path})
            _emit_step(
                on_step,
                {
                    "status": "excel_action_completed",
                    "agent": agent_label,
                    "action": action_name,
                    "result": "created",
                },
            )
            continue

        if action_name == "delete_workbook":
            if os.path.exists(workbook_path):
                os.remove(workbook_path)
                operation_logs.append({"action": action_name, "status": "ok", "workbook_path": workbook_path})
                _emit_step(
                    on_step,
                    {
                        "status": "excel_action_completed",
                        "agent": agent_label,
                        "action": action_name,
                        "result": "deleted",
                    },
                )
            else:
                operation_logs.append(
                    {
                        "action": action_name,
                        "status": "skipped",
                        "reason": "Workbook does not exist.",
                        "workbook_path": workbook_path,
                    }
                )
            continue

        if not os.path.exists(workbook_path):
            if action_auto_create_workbook:
                workbook = openpyxl.Workbook()
                workbook.active.title = default_sheet[:31] or "Sheet1"
                workbook.save(workbook_path)
                operation_logs.append(
                    {
                        "action": "auto_create_workbook",
                        "status": "ok",
                        "workbook_path": workbook_path,
                    }
                )
                _emit_step(
                    on_step,
                    {
                        "status": "excel_action_completed",
                        "agent": agent_label,
                        "action": "auto_create_workbook",
                        "result": "created_missing_workbook",
                    },
                )
            else:
                return {
                    "answer": (
                        f"Workbook not found at {workbook_path}. "
                        "Enable auto_create_workbook or create it first."
                    ),
                    "details": {
                        "action": action_name,
                        "workbook_path": workbook_path,
                        "auto_create_workbook": action_auto_create_workbook,
                    },
                }

        workbook = openpyxl.load_workbook(workbook_path)
        changed = False

        try:
            if action_name == "list_sheets":
                last_sheet_names = list(workbook.sheetnames)
                operation_logs.append({"action": action_name, "status": "ok", "sheets": last_sheet_names})
            elif action_name == "create_sheet":
                sheet_name = _resolve_excel_sheet_name_for_action(
                    workbook,
                    action,
                    default_sheet,
                    auto_sheet_selection,
                    action_name,
                )
                if sheet_name in workbook.sheetnames:
                    operation_logs.append(
                        {"action": action_name, "status": "skipped", "reason": "Sheet already exists.", "sheet": sheet_name}
                    )
                else:
                    workbook.create_sheet(title=sheet_name)
                    changed = True
                    operation_logs.append({"action": action_name, "status": "ok", "sheet": sheet_name})
            elif action_name == "delete_sheet":
                sheet_name = _resolve_excel_sheet_name_for_action(
                    workbook,
                    action,
                    default_sheet,
                    auto_sheet_selection,
                    action_name,
                )
                if sheet_name not in workbook.sheetnames:
                    operation_logs.append(
                        {"action": action_name, "status": "skipped", "reason": "Sheet not found.", "sheet": sheet_name}
                    )
                elif len(workbook.sheetnames) <= 1:
                    operation_logs.append(
                        {
                            "action": action_name,
                            "status": "skipped",
                            "reason": "Cannot delete the last remaining sheet.",
                            "sheet": sheet_name,
                        }
                    )
                else:
                    del workbook[sheet_name]
                    changed = True
                    operation_logs.append({"action": action_name, "status": "ok", "sheet": sheet_name})
            elif action_name == "rename_sheet":
                source = _resolve_excel_sheet_name_for_action(
                    workbook,
                    action,
                    default_sheet,
                    auto_sheet_selection,
                    action_name,
                )
                target = str(action.get("new_name") or "").strip()[:31]
                if source not in workbook.sheetnames:
                    operation_logs.append(
                        {"action": action_name, "status": "skipped", "reason": "Source sheet not found.", "sheet": source}
                    )
                elif not target:
                    operation_logs.append(
                        {"action": action_name, "status": "skipped", "reason": "Target sheet name is empty.", "sheet": source}
                    )
                elif target in workbook.sheetnames:
                    operation_logs.append(
                        {"action": action_name, "status": "skipped", "reason": "Target sheet already exists.", "sheet": target}
                    )
                else:
                    workbook[source].title = target
                    changed = True
                    operation_logs.append({"action": action_name, "status": "ok", "sheet": source, "new_name": target})
            elif action_name == "write_cells":
                sheet_name = _resolve_excel_sheet_name_for_action(
                    workbook,
                    action,
                    default_sheet,
                    auto_sheet_selection,
                    action_name,
                )
                if sheet_name not in workbook.sheetnames:
                    workbook.create_sheet(title=sheet_name[:31])
                    changed = True
                sheet = workbook[sheet_name[:31]]
                cells = action.get("cells")
                if not isinstance(cells, dict) or len(cells) == 0:
                    operation_logs.append(
                        {"action": action_name, "status": "skipped", "reason": "No cells payload provided.", "sheet": sheet_name}
                    )
                else:
                    written = 0
                    for cell_ref, value in cells.items():
                        ref = str(cell_ref).upper().strip()
                        if not re.fullmatch(r"[A-Z]{1,3}[1-9][0-9]{0,6}", ref):
                            continue
                        sheet[ref] = value
                        written += 1
                    changed = changed or (written > 0)
                    operation_logs.append(
                        {"action": action_name, "status": "ok", "sheet": sheet_name, "written_cells": written}
                    )
            elif action_name == "append_rows":
                sheet_name = _resolve_excel_sheet_name_for_action(
                    workbook,
                    action,
                    default_sheet,
                    auto_sheet_selection,
                    action_name,
                )
                if sheet_name not in workbook.sheetnames:
                    workbook.create_sheet(title=sheet_name[:31])
                    changed = True
                sheet = workbook[sheet_name[:31]]
                rows_payload = action.get("rows")
                if isinstance(rows_payload, list) and len(rows_payload) > 0:
                    appended = 0
                    for row in rows_payload:
                        if isinstance(row, list):
                            sheet.append(row)
                            appended += 1
                    changed = changed or (appended > 0)
                    operation_logs.append(
                        {"action": action_name, "status": "ok", "sheet": sheet_name, "appended_rows": appended}
                    )
                else:
                    operation_logs.append(
                        {"action": action_name, "status": "skipped", "reason": "No rows payload provided.", "sheet": sheet_name}
                    )
            elif action_name == "read_sheet":
                sheet_name = _resolve_excel_sheet_name_for_action(
                    workbook,
                    action,
                    default_sheet,
                    auto_sheet_selection,
                    action_name,
                )
                if sheet_name not in workbook.sheetnames:
                    operation_logs.append(
                        {"action": action_name, "status": "skipped", "reason": "Sheet not found.", "sheet": sheet_name}
                    )
                else:
                    limit = int(action.get("max_rows") or max_rows_read)
                    limit = max(1, min(limit, 10000))
                    sheet = workbook[sheet_name]
                    last_read_rows = []
                    for row in sheet.iter_rows(min_row=1, max_row=limit, values_only=True):
                        last_read_rows.append(list(row))
                    operation_logs.append(
                        {"action": action_name, "status": "ok", "sheet": sheet_name, "rows_read": len(last_read_rows)}
                    )
            else:
                operation_logs.append({"action": action_name, "status": "skipped", "reason": "Unsupported action."})
        finally:
            if changed:
                workbook.save(workbook_path)
            workbook.close()

        _emit_step(
            on_step,
            {
                "status": "excel_action_completed",
                "agent": agent_label,
                "action": action_name,
            },
        )

    summary_parts = [f"{len(operation_logs)} action(s) processed"]
    if workbook_path_cache:
        summary_parts.append(f"workbook={workbook_path_cache}")
    if last_sheet_names:
        summary_parts.append(f"sheets={', '.join(last_sheet_names)}")
    if last_read_rows:
        summary_parts.append(f"rows_read={len(last_read_rows)}")

    return {
        "answer": "Excel operations completed: " + " | ".join(summary_parts),
        "details": {
            "workbook_path": workbook_path_cache,
            "operations": operation_logs,
            "sheet_names": last_sheet_names,
            "rows": last_read_rows[:200],
        },
    }


def _is_instructional_word_document_value(value: str) -> bool:
    raw = str(value or "").strip()
    if not raw:
        return False
    lowered = raw.lower()
    if re.search(r"\.(docx|docm|dotx|dotm)$", lowered):
        return False
    if "/" in raw or "\\" in raw:
        return False

    word_count = len(re.findall(r"[A-Za-zÀ-ÿ0-9_'-]+", raw))
    instruction_markers = (
        "je veux",
        "i want",
        "ia decide",
        "ai decide",
        "auto",
        "folder",
        "dossier",
        "reste",
        "remain",
        "designe",
        "designated",
    )
    return word_count >= 6 and (len(raw) > 40 or any(marker in lowered for marker in instruction_markers))


def _normalize_word_document_preference(value: str) -> str:
    raw = str(value or "").strip()
    if not raw:
        return ""
    if _is_instructional_word_document_value(raw):
        return ""
    return raw


def _word_manager_uses_auto_document_selection(config: Dict[str, Any]) -> bool:
    mode = str(config.get("document_path_mode") or "").strip().lower()
    if mode in {"auto", "auto_in_folder", "dynamic", "llm_decides", "ia_decide"}:
        return True

    configured_document = str(config.get("document_path") or "").strip()
    return _is_instructional_word_document_value(configured_document)


def _list_word_documents_in_root(root_path: str) -> List[Path]:
    root = Path(root_path)
    suffixes = {".docx", ".docm", ".dotx", ".dotm"}
    documents: List[Path] = []
    try:
        for item in root.iterdir():
            if not item.is_file():
                continue
            if item.suffix.lower() not in suffixes:
                continue
            documents.append(item)
    except Exception:
        return []

    documents.sort(key=lambda path: path.stat().st_mtime, reverse=True)
    return documents


def _choose_auto_word_document_candidate(
    root_path: str,
    action_name: str,
    allow_overwrite: bool,
) -> str:
    existing_documents = _list_word_documents_in_root(root_path)
    normalized_action = str(action_name or "").strip().lower()

    if normalized_action in {"read_document", "append_paragraphs", "replace_text", "delete_document"}:
        if existing_documents:
            return str(existing_documents[0])
        return str(Path(root_path) / "report.docx")

    if normalized_action in {"create_document", "write_paragraphs"}:
        if existing_documents and allow_overwrite:
            return str(existing_documents[0])
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return str(Path(root_path) / f"document_{timestamp}.docx")

    if existing_documents:
        return str(existing_documents[0])
    return str(Path(root_path) / "report.docx")


def _resolve_word_document_path(
    config: Dict[str, Any],
    document_path: str | None = None,
    *,
    action_name: str | None = None,
) -> str:
    configured_document = _normalize_word_document_preference(str(config.get("document_path") or ""))
    requested_document = _normalize_word_document_preference(str(document_path or ""))
    root_path = _resolve_managed_root_path(config, default_subdir="data/documents")

    auto_document_selection = _word_manager_uses_auto_document_selection(config)
    if requested_document:
        raw_path = requested_document
    elif auto_document_selection:
        raw_path = _choose_auto_word_document_candidate(
            root_path=root_path,
            action_name=str(action_name or ""),
            allow_overwrite=bool(config.get("allow_overwrite")),
        )
    else:
        raw_path = configured_document or "report.docx"

    resolved = _resolve_safe_managed_file_path(
        root_path=root_path,
        raw_path=raw_path,
        allowed_suffixes=(".docx", ".docm", ".dotx", ".dotm"),
        allow_outside_folder=bool(config.get("allow_outside_folder")),
    )
    parent = os.path.dirname(resolved) or root_path
    if bool(config.get("auto_create_folder", True)):
        os.makedirs(parent, exist_ok=True)
    return resolved


def _guess_word_document_path_from_text(text: str) -> str | None:
    quoted = re.search(
        r"[\"']([^\"']+\.(?:docx|docm|dotx|dotm))[\"']",
        text or "",
        flags=re.IGNORECASE,
    )
    if quoted:
        return str(quoted.group(1)).strip()

    bare = re.search(
        r"([A-Za-z]:[\\/][^\s\"']+\.(?:docx|docm|dotx|dotm)|\.{0,2}[\\/][^\s\"']+\.(?:docx|docm|dotx|dotm)|[A-Za-z0-9._/-]+\.(?:docx|docm|dotx|dotm))",
        text or "",
        flags=re.IGNORECASE,
    )
    if bare:
        return str(bare.group(1)).strip()
    return None


def _coerce_word_paragraphs(value: Any) -> List[str]:
    if isinstance(value, str):
        lines = [line.strip() for line in value.splitlines()]
        return [line for line in lines if line]
    if isinstance(value, list):
        paragraphs: List[str] = []
        for item in value:
            text = str(item or "").strip()
            if text:
                paragraphs.append(text)
        return paragraphs
    return []


def _guess_replace_text_from_text(text: str) -> tuple[str, str] | None:
    patterns = [
        r"(?:replace|remplace(?:r)?)\s+[\"']([^\"']+)[\"']\s+(?:with|par|to)\s+[\"']([^\"']*)[\"']",
        r"(?:replace|remplace(?:r)?)(?:\s+the)?(?:\s+text)?\s+[\"']([^\"']+)[\"']\s+(?:with|par|to)\s+[\"']([^\"']*)[\"']",
    ]
    for pattern in patterns:
        match = re.search(pattern, text or "", flags=re.IGNORECASE)
        if match:
            old_text = str(match.group(1)).strip()
            new_text = str(match.group(2))
            if old_text:
                return old_text, new_text
    return None


def _guess_word_payload_from_text(text: str) -> List[str]:
    quoted = [str(item).strip() for item in re.findall(r"[\"']([^\"']+)[\"']", text or "")]
    quoted = [item for item in quoted if item]
    if quoted:
        return quoted

    colon_match = re.search(r"[:：]\s*(.+)$", text or "", flags=re.IGNORECASE)
    if colon_match:
        payload = str(colon_match.group(1)).strip()
        if payload:
            return [payload]
    return []


def _normalize_word_action_payload(action: Dict[str, Any]) -> Dict[str, Any]:
    normalized: Dict[str, Any] = dict(action)
    raw_action_name = str(normalized.get("action") or "").strip().lower()
    action_aliases = {
        "create_file": "create_document",
        "new_document": "create_document",
        "new_doc": "create_document",
        "delete_file": "delete_document",
        "remove_document": "delete_document",
        "remove_doc": "delete_document",
        "read": "read_document",
        "write": "write_paragraphs",
        "overwrite": "write_paragraphs",
        "append": "append_paragraphs",
        "append_text": "append_paragraphs",
        "replace": "replace_text",
    }
    normalized["action"] = action_aliases.get(raw_action_name, raw_action_name)

    if normalized.get("document_path") in (None, ""):
        for key in ("file_path", "file", "path", "document", "doc_path", "document_file"):
            value = normalized.get(key)
            if isinstance(value, str) and value.strip():
                normalized["document_path"] = value.strip()
                break

    action_name = str(normalized.get("action") or "")

    if action_name in {"write_paragraphs", "append_paragraphs"}:
        raw_paragraphs = normalized.get("paragraphs")
        if raw_paragraphs in (None, ""):
            for key in ("text", "content", "body", "lines", "paragraph"):
                if normalized.get(key) not in (None, ""):
                    raw_paragraphs = normalized.get(key)
                    break
        normalized["paragraphs"] = _coerce_word_paragraphs(raw_paragraphs)

    if action_name == "replace_text":
        old_text = normalized.get("old_text")
        new_text = normalized.get("new_text")
        if old_text in (None, ""):
            for key in ("from", "find", "target", "search", "old"):
                if normalized.get(key) not in (None, ""):
                    old_text = normalized.get(key)
                    break
        if new_text in (None, ""):
            for key in ("to", "replace", "replacement", "new"):
                if normalized.get(key) is not None:
                    new_text = normalized.get(key)
                    break
        normalized["old_text"] = str(old_text or "")
        normalized["new_text"] = str(new_text or "")

    if action_name == "read_document" and normalized.get("max_paragraphs") in (None, ""):
        for key in ("limit", "max_rows", "max_lines"):
            value = normalized.get(key)
            if value not in (None, ""):
                normalized["max_paragraphs"] = value
                break

    return normalized


def _normalize_word_actions(
    user_input: str,
    config: Dict[str, Any],
    allow_heuristics: bool = True,
) -> List[Dict[str, Any]] | None:
    parsed = _extract_possible_json(user_input)
    if isinstance(parsed, dict):
        if isinstance(parsed.get("actions"), list):
            actions = [_normalize_word_action_payload(action) for action in parsed["actions"] if isinstance(action, dict)]
            return actions if actions else None
        if isinstance(parsed.get("action"), str):
            return [_normalize_word_action_payload(parsed)]
    if isinstance(parsed, list):
        actions = [_normalize_word_action_payload(action) for action in parsed if isinstance(action, dict)]
        if actions:
            return actions

    if not allow_heuristics:
        return None

    text = (user_input or "").strip()
    lowered = text.lower()
    document_hint = _guess_word_document_path_from_text(text)
    max_paragraphs_read = int(config.get("max_paragraphs_read") or 100)
    heuristics: List[Dict[str, Any]] = []

    def add_action(action: Dict[str, Any]) -> None:
        if document_hint and not action.get("document_path"):
            action["document_path"] = document_hint
        heuristics.append(_normalize_word_action_payload(action))

    if re.search(r"(create|new|cr[eé]e|créer|add|ajoute).*(document|word|docx|fichier)", lowered):
        add_action({"action": "create_document"})
    elif re.search(r"(delete|remove|supprime|supprimer).*(document|word|docx|fichier)", lowered):
        add_action({"action": "delete_document"})
    elif re.search(r"(read|lire|show|affiche).*(document|word|docx|fichier)", lowered):
        add_action({"action": "read_document", "max_paragraphs": max_paragraphs_read})
    else:
        replace_payload = _guess_replace_text_from_text(text)
        if replace_payload:
            old_text, new_text = replace_payload
            add_action({"action": "replace_text", "old_text": old_text, "new_text": new_text})
        elif re.search(r"(append|add|ajoute).*(paragraph|paragraphe|texte)", lowered):
            paragraphs = _guess_word_payload_from_text(text)
            if paragraphs:
                add_action({"action": "append_paragraphs", "paragraphs": paragraphs})
        elif re.search(r"(write|ecris|écris|r[eé]dige|overwrite|replace all|remplace le contenu|modifie|modifier|edit)", lowered):
            paragraphs = _guess_word_payload_from_text(text)
            if paragraphs:
                add_action({"action": "write_paragraphs", "paragraphs": paragraphs})

    return heuristics if heuristics else None


def _clear_word_document_paragraphs(document: Any) -> None:
    for paragraph in list(document.paragraphs):
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


async def _execute_word_actions(
    user_input: str,
    config: Dict[str, Any],
    on_step: StepCallback | None,
    agent_label: str,
    allow_heuristics: bool = True,
) -> AgentResult | None:
    actions = _normalize_word_actions(user_input, config, allow_heuristics=allow_heuristics)
    if not actions:
        return None

    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        return {
            "answer": (
                "Word manager requires the `python-docx` dependency. "
                f"Install it to enable document operations. ({exc})"
            )
        }

    max_paragraphs_read = int(config.get("max_paragraphs_read") or 100)
    auto_create_document = bool(config.get("auto_create_document"))
    allow_overwrite = bool(config.get("allow_overwrite"))

    operation_logs: List[Dict[str, Any]] = []
    last_read_paragraphs: List[str] = []
    document_path_cache: str | None = None

    for action in actions:
        action_name = str(action.get("action") or "").strip().lower()
        try:
            document_path = _resolve_word_document_path(
                config,
                action.get("document_path"),
                action_name=action_name,
            )
        except Exception as exc:
            return {
                "answer": f"Unable to resolve document path: {exc}",
                "details": {"action": action_name, "requested_document_path": action.get("document_path")},
            }
        document_path_cache = document_path

        _emit_step(
            on_step,
            {
                "status": "word_action_started",
                "agent": agent_label,
                "action": action_name,
                "document_path": document_path,
            },
        )

        if action_name == "create_document":
            if os.path.exists(document_path) and not allow_overwrite:
                operation_logs.append(
                    {
                        "action": action_name,
                        "status": "skipped",
                        "reason": "Document already exists and overwrite is disabled.",
                        "document_path": document_path,
                    }
                )
                _emit_step(
                    on_step,
                    {
                        "status": "word_action_completed",
                        "agent": agent_label,
                        "action": action_name,
                        "result": "skipped_existing",
                    },
                )
                continue

            document = Document()
            paragraphs = _coerce_word_paragraphs(action.get("paragraphs"))
            for paragraph in paragraphs:
                document.add_paragraph(paragraph)
            document.save(document_path)
            operation_logs.append(
                {
                    "action": action_name,
                    "status": "ok",
                    "document_path": document_path,
                    "paragraphs_written": len(paragraphs),
                }
            )
            _emit_step(
                on_step,
                {
                    "status": "word_action_completed",
                    "agent": agent_label,
                    "action": action_name,
                    "result": "created",
                },
            )
            continue

        if action_name == "delete_document":
            if os.path.exists(document_path):
                os.remove(document_path)
                operation_logs.append({"action": action_name, "status": "ok", "document_path": document_path})
                _emit_step(
                    on_step,
                    {
                        "status": "word_action_completed",
                        "agent": agent_label,
                        "action": action_name,
                        "result": "deleted",
                    },
                )
            else:
                operation_logs.append(
                    {
                        "action": action_name,
                        "status": "skipped",
                        "reason": "Document does not exist.",
                        "document_path": document_path,
                    }
                )
            continue

        if not os.path.exists(document_path):
            if auto_create_document:
                document = Document()
                document.save(document_path)
            else:
                return {
                    "answer": (
                        f"Document not found at {document_path}. "
                        "Enable auto_create_document or create it first."
                    ),
                    "details": {
                        "action": action_name,
                        "document_path": document_path,
                    },
                }

        try:
            document = Document(document_path)
        except Exception as exc:
            return {
                "answer": f"Unable to open document at {document_path}: {exc}",
                "details": {"action": action_name, "document_path": document_path},
            }
        changed = False

        if action_name == "read_document":
            limit = int(action.get("max_paragraphs") or max_paragraphs_read)
            limit = max(1, min(limit, 5000))
            last_read_paragraphs = []
            for paragraph in document.paragraphs:
                text = str(paragraph.text or "").strip()
                if not text:
                    continue
                last_read_paragraphs.append(text)
                if len(last_read_paragraphs) >= limit:
                    break
            operation_logs.append(
                {
                    "action": action_name,
                    "status": "ok",
                    "paragraphs_read": len(last_read_paragraphs),
                }
            )
        elif action_name == "write_paragraphs":
            paragraphs = _coerce_word_paragraphs(action.get("paragraphs"))
            if len(paragraphs) == 0:
                operation_logs.append(
                    {"action": action_name, "status": "skipped", "reason": "No paragraph payload provided."}
                )
            else:
                _clear_word_document_paragraphs(document)
                for paragraph in paragraphs:
                    document.add_paragraph(paragraph)
                changed = True
                operation_logs.append(
                    {
                        "action": action_name,
                        "status": "ok",
                        "paragraphs_written": len(paragraphs),
                    }
                )
        elif action_name == "append_paragraphs":
            paragraphs = _coerce_word_paragraphs(action.get("paragraphs"))
            if len(paragraphs) == 0:
                operation_logs.append(
                    {"action": action_name, "status": "skipped", "reason": "No paragraph payload provided."}
                )
            else:
                for paragraph in paragraphs:
                    document.add_paragraph(paragraph)
                changed = True
                operation_logs.append(
                    {
                        "action": action_name,
                        "status": "ok",
                        "paragraphs_appended": len(paragraphs),
                    }
                )
        elif action_name == "replace_text":
            old_text = str(action.get("old_text") or "")
            new_text = str(action.get("new_text") or "")
            if not old_text:
                operation_logs.append(
                    {"action": action_name, "status": "skipped", "reason": "No source text provided."}
                )
            else:
                replaced_count = 0
                for paragraph in document.paragraphs:
                    content = str(paragraph.text or "")
                    if old_text in content:
                        paragraph.text = content.replace(old_text, new_text)
                        replaced_count += 1
                if replaced_count > 0:
                    changed = True
                    operation_logs.append(
                        {
                            "action": action_name,
                            "status": "ok",
                            "paragraphs_replaced": replaced_count,
                        }
                    )
                else:
                    operation_logs.append(
                        {"action": action_name, "status": "skipped", "reason": "Text not found in document."}
                    )
        else:
            operation_logs.append({"action": action_name, "status": "skipped", "reason": "Unsupported action."})

        if changed:
            document.save(document_path)

        _emit_step(
            on_step,
            {
                "status": "word_action_completed",
                "agent": agent_label,
                "action": action_name,
            },
        )

    summary_parts = [f"{len(operation_logs)} action(s) processed"]
    if document_path_cache:
        summary_parts.append(f"document={document_path_cache}")
    if last_read_paragraphs:
        summary_parts.append(f"paragraphs_read={len(last_read_paragraphs)}")

    return {
        "answer": "Word operations completed: " + " | ".join(summary_parts),
        "details": {
            "document_path": document_path_cache,
            "operations": operation_logs,
            "paragraphs": last_read_paragraphs[:500],
        },
    }


def _coerce_bool(value: Any, default: bool = False) -> bool:
    if isinstance(value, bool):
        return value
    if value is None:
        return default
    lowered = str(value).strip().lower()
    if lowered in {"1", "true", "yes", "on", "y"}:
        return True
    if lowered in {"0", "false", "no", "off", "n"}:
        return False
    return default


def _coerce_int(value: Any, default: int, minimum: int, maximum: int) -> int:
    try:
        parsed = int(value)
    except Exception:
        parsed = default
    return max(minimum, min(maximum, parsed))


def _looks_like_file_path(value: str) -> bool:
    text = str(value or "").strip()
    if not text:
        return False
    if "/" in text or "\\" in text:
        return True
    return bool(re.search(r"\.[A-Za-z0-9]{1,8}$", text))


def _resolve_text_file_path(config: Dict[str, Any], file_path: str | None = None) -> tuple[str, str]:
    root_path = _resolve_managed_root_path(config, default_subdir="data/workspace")
    configured_path = str(config.get("default_file_path") or "").strip()
    requested_path = str(file_path or "").strip()
    raw_path = requested_path or configured_path or "output.txt"
    resolved = _resolve_safe_managed_file_path(
        root_path=root_path,
        raw_path=raw_path,
        allow_outside_folder=_coerce_bool(config.get("allow_outside_folder"), default=False),
    )
    return root_path, resolved


def _guess_text_file_path_from_text(text: str) -> str | None:
    quoted = re.search(
        r"[\"']([^\"']+\.(?:txt|md|csv|json|log|yaml|yml|ini|cfg|conf|xml|html|py|js|ts))[\"']",
        text or "",
        flags=re.IGNORECASE,
    )
    if quoted:
        return str(quoted.group(1)).strip()

    bare = re.search(
        r"([A-Za-z]:[\\/][^\s\"']+\.(?:txt|md|csv|json|log|yaml|yml|ini|cfg|conf|xml|html|py|js|ts)|\.{0,2}[\\/][^\s\"']+\.(?:txt|md|csv|json|log|yaml|yml|ini|cfg|conf|xml|html|py|js|ts)|[A-Za-z0-9._/-]+\.(?:txt|md|csv|json|log|yaml|yml|ini|cfg|conf|xml|html|py|js|ts))",
        text or "",
        flags=re.IGNORECASE,
    )
    if bare:
        return str(bare.group(1)).strip()
    return None


def _guess_text_content_from_text(text: str) -> str:
    quoted_values = [str(item).strip() for item in re.findall(r"[\"']([^\"']+)[\"']", text or "")]
    filtered = [item for item in quoted_values if item and not _looks_like_file_path(item)]
    if filtered:
        return filtered[-1]

    colon_match = re.search(r"[:：]\s*(.+)$", text or "", flags=re.IGNORECASE)
    if colon_match:
        return str(colon_match.group(1)).strip()
    return ""


def _normalize_text_action_payload(action: Dict[str, Any]) -> Dict[str, Any]:
    normalized: Dict[str, Any] = dict(action)
    raw_action_name = str(normalized.get("action") or normalized.get("operation") or "").strip().lower()
    action_aliases = {
        "create": "create_file",
        "new_file": "create_file",
        "delete": "delete_file",
        "remove": "delete_file",
        "read": "read_file",
        "write": "write_file",
        "append": "append_file",
        "list": "list_files",
    }
    normalized["action"] = action_aliases.get(raw_action_name, raw_action_name)

    if normalized.get("file_path") in (None, ""):
        for key in ("file", "path", "target_path", "filename"):
            value = normalized.get(key)
            if isinstance(value, str) and value.strip():
                normalized["file_path"] = value.strip()
                break

    if normalized.get("content") is None:
        for key in ("text", "value", "body", "data"):
            if key in normalized and normalized.get(key) is not None:
                normalized["content"] = normalized.get(key)
                break

    if normalized.get("pattern") in (None, "") and normalized.get("filter") not in (None, ""):
        normalized["pattern"] = normalized.get("filter")

    return normalized


def _normalize_text_file_actions(
    user_input: str,
    config: Dict[str, Any],
    *,
    allow_heuristics: bool = True,
) -> List[Dict[str, Any]] | None:
    parsed = _extract_possible_json(user_input)
    if isinstance(parsed, dict):
        if isinstance(parsed.get("actions"), list):
            actions = [_normalize_text_action_payload(item) for item in parsed["actions"] if isinstance(item, dict)]
            return actions if actions else None
        if isinstance(parsed.get("action"), str) or isinstance(parsed.get("operation"), str):
            return [_normalize_text_action_payload(parsed)]
    if isinstance(parsed, list):
        actions = [_normalize_text_action_payload(item) for item in parsed if isinstance(item, dict)]
        if actions:
            return actions

    if not allow_heuristics:
        return None

    text = (user_input or "").strip()
    lowered = text.lower()
    guessed_path = _guess_text_file_path_from_text(text)
    guessed_content = _guess_text_content_from_text(text)

    heuristics: List[Dict[str, Any]] = []

    def add_action(action: Dict[str, Any]) -> None:
        if guessed_path and not action.get("file_path"):
            action["file_path"] = guessed_path
        heuristics.append(_normalize_text_action_payload(action))

    if re.search(r"(list|liste|show|affiche).*(files?|fichiers?)", lowered):
        add_action({"action": "list_files"})
    elif re.search(r"(delete|remove|supprime|supprimer).*(files?|fichiers?)", lowered):
        add_action({"action": "delete_file"})
    elif re.search(r"(create|new|cr[eé]e|créer).*(files?|fichiers?)", lowered):
        payload: Dict[str, Any] = {"action": "create_file"}
        if guessed_content:
            payload["content"] = guessed_content
        add_action(payload)
    elif re.search(r"(read|lire|open|ouvre|affiche).*(files?|fichiers?)", lowered):
        add_action({"action": "read_file", "max_chars": int(config.get("max_chars_read") or 10000)})
    elif re.search(r"(append|ajoute|ajouter).*(text|texte|files?|fichiers?)", lowered):
        if guessed_content:
            add_action({"action": "append_file", "content": guessed_content})
    elif re.search(r"(write|ecris|écris|modifie|modifier|replace).*(files?|fichiers?|text|texte)", lowered):
        if guessed_content:
            add_action({"action": "write_file", "content": guessed_content})

    return heuristics if heuristics else None


async def _execute_text_file_actions(
    user_input: str,
    config: Dict[str, Any],
    on_step: StepCallback | None,
    agent_label: str,
    *,
    allow_heuristics: bool = True,
) -> AgentResult | None:
    actions = _normalize_text_file_actions(user_input, config, allow_heuristics=allow_heuristics)
    if not actions:
        return None

    default_encoding = str(config.get("default_encoding") or "utf-8").strip() or "utf-8"
    allow_overwrite = _coerce_bool(config.get("allow_overwrite"), default=True)
    max_chars_read = _coerce_int(config.get("max_chars_read"), default=10000, minimum=200, maximum=500000)

    operation_logs: List[Dict[str, Any]] = []
    listed_files: List[Dict[str, Any]] = []
    preview_by_file: Dict[str, str] = {}
    root_path_cache = ""
    file_path_cache = ""

    for action in actions:
        action_name = str(action.get("action") or "").strip().lower()
        raw_file_path = str(action.get("file_path") or "").strip()
        resolved_file_path = ""

        try:
            if action_name == "list_files":
                root_path = _resolve_managed_root_path(config, default_subdir="data/workspace")
                root_path_cache = root_path
                target_path = root_path
                if raw_file_path:
                    target_path = _resolve_safe_managed_file_path(
                        root_path=root_path,
                        raw_path=raw_file_path,
                        allow_outside_folder=_coerce_bool(config.get("allow_outside_folder"), default=False),
                    )
                resolved_file_path = target_path
            else:
                root_path, resolved_file_path = _resolve_text_file_path(config, raw_file_path or None)
                root_path_cache = root_path
                file_path_cache = resolved_file_path
        except Exception as exc:
            return {
                "answer": f"Unable to resolve text-file path: {exc}",
                "details": {"action": action_name, "requested_file_path": raw_file_path},
            }

        _emit_step(
            on_step,
            {
                "status": "text_action_started",
                "agent": agent_label,
                "action": action_name,
                "file_path": resolved_file_path,
            },
        )

        if action_name == "list_files":
            recursive = _coerce_bool(action.get("recursive"), default=False)
            pattern = str(action.get("pattern") or "").strip().lower()
            entries: List[Dict[str, Any]] = []

            if os.path.isfile(resolved_file_path):
                entries = [
                    {
                        "path": _to_relative_path(root_path_cache, resolved_file_path),
                        "size_bytes": os.path.getsize(resolved_file_path),
                    }
                ]
            elif os.path.isdir(resolved_file_path):
                iterator = Path(resolved_file_path).rglob("*") if recursive else Path(resolved_file_path).glob("*")
                for item in iterator:
                    if not item.is_file():
                        continue
                    relative = _to_relative_path(root_path_cache, str(item))
                    if pattern and pattern not in relative.lower():
                        continue
                    entries.append({"path": relative, "size_bytes": item.stat().st_size})
                    if len(entries) >= 500:
                        break
            else:
                operation_logs.append(
                    {
                        "action": action_name,
                        "status": "skipped",
                        "reason": "Target path does not exist.",
                        "file_path": resolved_file_path,
                    }
                )
                continue

            listed_files = entries
            operation_logs.append(
                {
                    "action": action_name,
                    "status": "ok",
                    "listed_files": len(entries),
                    "target": _to_relative_path(root_path_cache, resolved_file_path),
                }
            )
            _emit_step(
                on_step,
                {
                    "status": "text_action_completed",
                    "agent": agent_label,
                    "action": action_name,
                    "listed_files": len(entries),
                },
            )
            continue

        if action_name == "read_file":
            if not os.path.exists(resolved_file_path) or not os.path.isfile(resolved_file_path):
                operation_logs.append(
                    {
                        "action": action_name,
                        "status": "skipped",
                        "reason": "File does not exist.",
                        "file_path": resolved_file_path,
                    }
                )
                continue

            max_chars = _coerce_int(action.get("max_chars"), default=max_chars_read, minimum=200, maximum=500000)
            content = Path(resolved_file_path).read_text(encoding=default_encoding, errors="ignore")
            preview = content[:max_chars]
            if len(content) > max_chars:
                preview += "\n...[truncated]"
            relative = _to_relative_path(root_path_cache, resolved_file_path)
            preview_by_file[relative] = preview
            operation_logs.append(
                {
                    "action": action_name,
                    "status": "ok",
                    "file_path": relative,
                    "chars_read": len(content),
                }
            )
            _emit_step(
                on_step,
                {
                    "status": "text_action_completed",
                    "agent": agent_label,
                    "action": action_name,
                    "file_path": relative,
                },
            )
            continue

        content = str(action.get("content") or "")
        target = Path(resolved_file_path)
        target.parent.mkdir(parents=True, exist_ok=True)
        relative_path = _to_relative_path(root_path_cache, resolved_file_path)

        if action_name == "create_file":
            if target.exists() and not allow_overwrite:
                operation_logs.append(
                    {
                        "action": action_name,
                        "status": "skipped",
                        "reason": "File already exists and overwrite is disabled.",
                        "file_path": relative_path,
                    }
                )
            else:
                if content:
                    target.write_text(content, encoding=default_encoding)
                else:
                    target.touch(exist_ok=True)
                operation_logs.append(
                    {
                        "action": action_name,
                        "status": "ok",
                        "file_path": relative_path,
                        "chars_written": len(content),
                    }
                )
        elif action_name == "write_file":
            if target.exists() and not allow_overwrite:
                operation_logs.append(
                    {
                        "action": action_name,
                        "status": "skipped",
                        "reason": "Overwrite disabled.",
                        "file_path": relative_path,
                    }
                )
            else:
                target.write_text(content, encoding=default_encoding)
                operation_logs.append(
                    {
                        "action": action_name,
                        "status": "ok",
                        "file_path": relative_path,
                        "chars_written": len(content),
                    }
                )
        elif action_name == "append_file":
            prefix = ""
            if _coerce_bool(action.get("prepend_newline"), default=False) and target.exists() and target.stat().st_size > 0:
                prefix = "\n"
            with target.open("a", encoding=default_encoding) as handle:
                handle.write(prefix + content)
            operation_logs.append(
                {
                    "action": action_name,
                    "status": "ok",
                    "file_path": relative_path,
                    "chars_appended": len(prefix + content),
                }
            )
        elif action_name == "delete_file":
            if target.exists():
                target.unlink()
                operation_logs.append({"action": action_name, "status": "ok", "file_path": relative_path})
            else:
                operation_logs.append(
                    {
                        "action": action_name,
                        "status": "skipped",
                        "reason": "File does not exist.",
                        "file_path": relative_path,
                    }
                )
        else:
            operation_logs.append({"action": action_name, "status": "skipped", "reason": "Unsupported action."})

        _emit_step(
            on_step,
            {
                "status": "text_action_completed",
                "agent": agent_label,
                "action": action_name,
                "file_path": relative_path,
            },
        )

    summary_parts = [f"{len(operation_logs)} action(s) processed"]
    if root_path_cache:
        summary_parts.append(f"folder={root_path_cache}")
    if file_path_cache:
        summary_parts.append(f"file={file_path_cache}")
    if listed_files:
        summary_parts.append(f"listed_files={len(listed_files)}")
    if preview_by_file:
        summary_parts.append(f"files_read={len(preview_by_file)}")

    return {
        "answer": "Text file operations completed: " + " | ".join(summary_parts),
        "details": {
            "folder_path": root_path_cache,
            "file_path": file_path_cache,
            "operations": operation_logs,
            "listed_files": listed_files[:500],
            "previews": preview_by_file,
        },
    }


def build_agent_execution_config(agent: Dict[str, Any]) -> Dict[str, Any]:
    parsed: Dict[str, Any] = {}
    raw_config = agent.get("config")

    if isinstance(raw_config, str):
        try:
            parsed = json.loads(raw_config)
        except Exception:
            parsed = {}
    elif isinstance(raw_config, dict):
        parsed = raw_config

    return {
        **parsed,
        "role": agent.get("role"),
        "persona": agent.get("persona"),
        "objectives": agent.get("objectives"),
        "system_prompt": agent.get("system_prompt"),
    }


def _is_generic_orchestration_text(value: str) -> bool:
    lowered = value.strip().lower()
    if not lowered:
        return True

    generic_markers = (
        "starting multi-agent orchestration",
        "response ready",
        "analyzing request and preparing context",
        "generating response with the language model",
        "model response received",
    )
    return any(marker in lowered for marker in generic_markers)


def _extract_manager_answer_from_steps(steps: List[Dict[str, Any]]) -> str:
    for step in reversed(steps):
        status = str(step.get("status") or "")

        if status in {"manager_final", "manager_finalized_fallback"}:
            for key in ("answer", "manager_summary"):
                value = step.get(key)
                if isinstance(value, str) and value.strip() and not _is_generic_orchestration_text(value):
                    return value.strip()

        if status in {"agent_call_completed"}:
            result = step.get("result")
            if isinstance(result, dict):
                value = result.get("answer")
                if isinstance(value, str) and value.strip() and not _is_generic_orchestration_text(value):
                    return value.strip()

        if status in {"agent_call_failed", "manager_finalize_failed", "error"}:
            for key in ("error", "message"):
                value = step.get(key)
                if isinstance(value, str) and value.strip() and not _is_generic_orchestration_text(value):
                    return value.strip()
    return ""


def _summarize_steps_for_synthesis(steps: List[Dict[str, Any]], max_steps: int = 25) -> str:
    lines: List[str] = []
    for step in steps[-max_steps:]:
        status = str(step.get("status") or "unknown")
        agent = str(step.get("agent") or "")
        message = str(step.get("message") or "")
        rationale = str(step.get("rationale") or "")
        result = step.get("result")
        result_text = ""
        if isinstance(result, dict):
            result_text = str(result.get("answer") or "")
        sql = str(step.get("sql") or "")

        parts = [f"status={status}"]
        if agent:
            parts.append(f"agent={agent}")
        if message:
            parts.append(f"message={message}")
        if rationale:
            parts.append(f"rationale={rationale}")
        if result_text:
            parts.append(f"result={result_text}")
        if sql:
            parts.append(f"sql={sql[:400]}")
        lines.append("- " + " | ".join(parts))

    return "\n".join(lines)


async def _synthesize_manager_fallback_answer(
    llm: Any,
    user_input: str,
    conversation_history: str,
    scratchpad: Dict[str, Any],
    steps: List[Dict[str, Any]],
) -> str:
    prompt = f"""You are finalizing a manager-orchestrated task.
Generate a complete, user-facing final answer based on the execution trace.
Do not mention internal implementation details.
If information is missing, state exactly what is missing.

Return ONLY JSON:
{{ "final_answer": "..." }}

User request:
{user_input}

Conversation history:
{conversation_history}

Scratchpad:
{json.dumps(scratchpad, ensure_ascii=False)}

Execution steps:
{_summarize_steps_for_synthesis(steps)}
"""

    response = await llm.ainvoke([HumanMessage(content=prompt)])
    content = _coerce_message_content(response.content).strip()
    if not content:
        return ""

    try:
        parsed = json.loads(_extract_json_block(content))
        if isinstance(parsed, dict):
            return str(parsed.get("final_answer") or parsed.get("answer") or "").strip()
    except Exception:
        pass

    return content


def _collect_param_assignments_from_text(text: str) -> Dict[str, Any]:
    assignments: Dict[str, Any] = {}
    pattern = re.compile(
        r"(?<![A-Za-z0-9_])([A-Za-z][A-Za-z0-9_]*)\s*(?:=|:)\s*(\"[^\"]*\"|'[^']*'|[^,\n;]+)"
    )
    for match in pattern.finditer(text or ""):
        key = _normalize_clickhouse_specific_param_name(match.group(1))
        if not key:
            continue
        raw_value = str(match.group(2)).strip().rstrip(".")
        if (raw_value.startswith('"') and raw_value.endswith('"')) or (
            raw_value.startswith("'") and raw_value.endswith("'")
        ):
            raw_value = raw_value[1:-1]
        assignments[key] = raw_value
    return assignments


def _normalize_clickhouse_specific_params(raw_params: Any) -> Dict[str, Any]:
    if not isinstance(raw_params, dict):
        return {}
    normalized: Dict[str, Any] = {}
    for key, value in raw_params.items():
        param_name = _normalize_clickhouse_specific_param_name(key)
        if not param_name:
            continue
        normalized[param_name] = value
    return normalized


def _resolve_clickhouse_specific_template(
    templates: List[Dict[str, Any]],
    query_name: str,
    default_query: str,
) -> Dict[str, Any] | None:
    if len(templates) == 0:
        return None

    requested = _normalize_identifier(query_name)
    if requested:
        for template in templates:
            if _normalize_identifier(template.get("name")) == requested:
                return template

    default_normalized = _normalize_identifier(default_query)
    if default_normalized:
        for template in templates:
            if _normalize_identifier(template.get("name")) == default_normalized:
                return template

    return templates[0]


def _looks_like_number(value: str) -> bool:
    return bool(re.fullmatch(r"-?\d+(\.\d+)?", value.strip()))


def _format_clickhouse_specific_value(value: Any, quote_mode: str) -> str:
    if value is None:
        return "NULL"

    if isinstance(value, list):
        return ", ".join(_format_clickhouse_specific_value(item, quote_mode) for item in value)

    if isinstance(value, bool):
        return "1" if value else "0"

    if isinstance(value, (int, float)):
        return str(value)

    text = str(value)
    normalized_quote_mode = (quote_mode or "auto").strip().lower()
    if normalized_quote_mode == "none":
        return text
    if normalized_quote_mode == "auto" and _looks_like_number(text):
        return text

    escaped = text.replace("'", "''")
    return f"'{escaped}'"


def _replace_clickhouse_specific_placeholder(sql: str, token: str, replacement: str) -> tuple[str, int]:
    replaced = 0
    output = sql

    output, count_braced = re.subn(r"\{" + re.escape(token) + r"\}", replacement, output)
    replaced += count_braced

    output, count_plain = re.subn(
        r"(?<![A-Za-z0-9_])" + re.escape(token) + r"(?![A-Za-z0-9_])",
        replacement,
        output,
    )
    replaced += count_plain

    return output, replaced


class AgentExecutor:
    @staticmethod
    async def execute(
        agent_type: str,
        config: Dict[str, Any],
        user_input: str,
        llm: Any,
        on_step: StepCallback | None = None,
        agent_name: str | None = None,
        invocation: Dict[str, Any] | None = None,
    ) -> AgentResult:
        system_prompt = "You are a helpful AI assistant."
        extra_context = ""
        agent_label = agent_name or agent_type or "agent"

        _emit_step(
            on_step,
            {
                "status": "agent_thinking",
                "agent": agent_label,
                "message": "Analyzing request and preparing context.",
            },
        )

        if agent_type == "sql_analyst":
            system_prompt = f"""You are a Senior ClickHouse Expert and SQL Analyst. Database: {config.get('database_name') or 'Unknown'}. 
Mode: {config.get('sql_use_case_mode') or 'llm_sql'}.
Template: {config.get('sql_query_template') or 'None'}.

CRITICAL CLICKHOUSE BEST PRACTICES:
- Never use PostgreSQL or MySQL syntax.
- For time-series gap analysis, use WITH FILL.
- For behavioral analysis (funnels), use windowFunnel().
- For fast descriptive stats on billions of rows, use quantileTiming, uniqHLL12, topK.
- For conditional aggregations, prefer sumIf, countIf over complex joins.
- For time-series joins, always consider ASOF JOIN.
- For distinct counts on large volumes, use uniq or uniqExact.
- Filters MUST use the table's sorting key (ORDER BY) for optimal performance.

CRITICAL FOR LARGE SCHEMAS & HUGE TABLES:
- Do not guess table or column names.
- If you lack schema context, write exploratory SQL first (e.g., SHOW TABLES, DESCRIBE table, or query system.columns/system.tables).
- Always LIMIT your queries when exploring data (e.g., SELECT * FROM table LIMIT 5).
- If parameters are missing, state what is missing in your answer.

Return your response in JSON format: {{ "sql": "SELECT ...", "answer": "Explanation..." }}"""
        elif agent_type == "clickhouse_table_manager":
            system_prompt = f"""You are a ClickHouse Table Manager. 
Protect existing tables: {config.get('protect_existing_tables') is not False}.
Allow inserts: {bool(config.get('allow_row_inserts'))}.
Allow updates: {bool(config.get('allow_row_updates'))}.
Allow deletes: {bool(config.get('allow_row_deletes'))}.

CRITICAL FOR LARGE SCHEMAS:
- Verify table existence and structure before applying changes.
- Do not drop or truncate tables unless explicitly permitted and verified.

Return your response in JSON format: {{ "sql": "CREATE TABLE ...", "answer": "Table created." }}"""
        elif agent_type == "clickhouse_writer":
            system_prompt = f"""You are a ClickHouse Data Writer.
Enforce agent prefix: {config.get('enforce_agent_prefix') is not False}.
Allow inserts: {config.get('allow_inserts') is not False}.

CRITICAL SECURITY RULE:
- You are ONLY allowed to create, modify, or insert into tables whose names start with "agent_".
- If the user asks to modify a table that does not start with "agent_", you MUST refuse and explain the security policy.
- Ensure all inserted data is properly escaped and formatted.

Return your response in JSON format: {{ "sql": "INSERT INTO agent_... ", "answer": "Data inserted." }}"""
        elif agent_type == "clickhouse_specific":
            templates = _normalize_clickhouse_specific_templates(config)
            default_query = str(config.get("default_query") or "").strip()

            invocation_payload = invocation if isinstance(invocation, dict) else {}
            query_name = str(
                invocation_payload.get("query_name")
                or invocation_payload.get("template_name")
                or ""
            ).strip()

            provided_params: Dict[str, Any] = {}
            provided_params.update(_normalize_clickhouse_specific_params(invocation_payload.get("params")))
            provided_params.update(_normalize_clickhouse_specific_params(invocation_payload.get("parameters")))

            parsed_input = _extract_possible_json(user_input)
            effective_request = user_input
            if isinstance(parsed_input, dict):
                query_name = query_name or str(
                    parsed_input.get("query_name")
                    or parsed_input.get("template_name")
                    or parsed_input.get("query")
                    or ""
                ).strip()
                provided_params.update(_normalize_clickhouse_specific_params(parsed_input.get("params")))
                provided_params.update(_normalize_clickhouse_specific_params(parsed_input.get("parameters")))
                for key in ("input", "request", "message", "task"):
                    value = parsed_input.get(key)
                    if isinstance(value, str) and value.strip():
                        effective_request = value.strip()
                        break

            text_assignments = _collect_param_assignments_from_text(effective_request)
            for key, value in text_assignments.items():
                provided_params.setdefault(key, value)

            selected_template = _resolve_clickhouse_specific_template(templates, query_name, default_query)
            if selected_template is None:
                available = [template.get("name") for template in templates if template.get("name")]
                message = (
                    "No ClickHouse query template is configured for this agent. "
                    "Add `query_templates` in agent config."
                )
                if available:
                    message += f" Available templates: {', '.join(str(item) for item in available)}."
                _emit_step(
                    on_step,
                    {
                        "status": "agent_completed",
                        "agent": agent_label,
                        "message": message,
                    },
                )
                return {
                    "answer": message,
                    "details": {
                        "available_templates": available,
                    },
                }

            sql_template = str(selected_template.get("sql") or "").strip()
            template_name = str(selected_template.get("name") or "default").strip() or "default"
            param_specs = selected_template.get("parameters")
            if not isinstance(param_specs, list):
                param_specs = []

            normalized_lookup = {
                _normalize_identifier(key): value for key, value in provided_params.items()
            }

            rendered_sql = sql_template
            applied_params: Dict[str, Any] = {}
            missing_params: List[str] = []

            for raw_spec in param_specs:
                if not isinstance(raw_spec, dict):
                    continue
                param_name = _normalize_clickhouse_specific_param_name(raw_spec.get("name"))
                if not param_name:
                    continue

                normalized_param_name = _normalize_identifier(param_name)
                value = normalized_lookup.get(normalized_param_name)
                if value in (None, ""):
                    default_value = raw_spec.get("default")
                    if default_value not in (None, ""):
                        value = default_value

                required = bool(raw_spec.get("required", True))
                if value in (None, ""):
                    if required:
                        missing_params.append(param_name)
                    continue

                replacement = _format_clickhouse_specific_value(value, str(raw_spec.get("quote") or "auto"))
                rendered_sql, _ = _replace_clickhouse_specific_placeholder(rendered_sql, param_name, replacement)
                applied_params[param_name] = value

            unresolved_tokens = [
                token for token in _extract_clickhouse_specific_tokens_from_sql(rendered_sql)
                if _normalize_identifier(token) in {
                    _normalize_identifier(str(spec.get("name") or ""))
                    for spec in param_specs
                    if isinstance(spec, dict)
                }
            ]
            for unresolved in unresolved_tokens:
                if unresolved not in missing_params:
                    missing_params.append(unresolved)

            if len(missing_params) > 0:
                missing_unique = list(dict.fromkeys(missing_params))
                missing_label = ", ".join(missing_unique)
                answer = (
                    f"Template `{template_name}` selected but missing required parameters: {missing_label}. "
                    "Provide them as `P1=value` in message or via manager call `params`."
                )
                _emit_step(
                    on_step,
                    {
                        "status": "agent_completed",
                        "agent": agent_label,
                        "message": "Missing required parameters for ClickHouse template.",
                        "template_name": template_name,
                        "missing_params": missing_unique,
                    },
                )
                return {
                    "answer": answer,
                    "sql": sql_template,
                    "details": {
                        "template_name": template_name,
                        "missing_params": missing_unique,
                        "provided_params": provided_params,
                        "available_templates": [
                            template.get("name") for template in templates if template.get("name")
                        ],
                    },
                }

            _emit_step(
                on_step,
                {
                    "status": "sql_template_rendered",
                    "agent": agent_label,
                    "template_name": template_name,
                    "params": applied_params,
                    "sql": rendered_sql,
                },
            )
            _emit_step(
                on_step,
                {
                    "status": "agent_completed",
                    "agent": agent_label,
                    "message": "Parameterized SQL rendered successfully.",
                },
            )
            return {
                "answer": (
                    f"Template `{template_name}` prepared successfully with "
                    f"{len(applied_params)} parameter(s)."
                ),
                "sql": rendered_sql,
                "details": {
                    "template_name": template_name,
                    "applied_params": applied_params,
                    "sql_template": sql_template,
                },
            }
        elif agent_type == "knowledge_base_assistant":
            system_prompt = f"""You are an internal Knowledge Base Assistant (Concierge Interne).
Vector DB: {config.get('vector_db') or 'qdrant'}
Embedding Model: {config.get('embedding_model') or 'text-embedding-004'}
Use Hybrid Search: {bool(config.get('use_hybrid_search'))}
Use Reranker: {bool(config.get('use_reranker'))}

CRITICAL RAG BEST PRACTICES:
- NEVER hallucinate. If the answer is not in the provided context, state clearly that you do not know.
- ALWAYS cite your sources (e.g., [Document Name, Page X]).
- Synthesize the retrieved chunks into a coherent, professional response.

Return your response in JSON format: {{ "answer": "Your synthesized answer with citations...", "sources": ["doc1", "doc2"] }}"""
        elif agent_type == "data_anomaly_hunter":
            system_prompt = f"""You are a Data Anomaly Hunter.
Detection Method: {config.get('detection_method') or 'statistical'}
Use Dynamic Thresholds: {bool(config.get('use_dynamic_thresholds'))}

CRITICAL ANOMALY DETECTION BEST PRACTICES:
- You receive statistical outliers (e.g., Z-score > 3, IQR anomalies) and KPI baselines.
- DO NOT read raw data rows. Rely on the statistical summaries provided.
- Explain WHY the anomaly might have happened in clear business terms.
- Format an alert message suitable for {config.get('alert_channel') or 'Slack'}.

Return your response in JSON format: {{ "answer": "Business explanation of the anomaly...", "alert_message": "🚨 *Anomaly Detected*..." }}"""
        elif agent_type == "text_to_sql_translator":
            system_prompt = f"""You are an expert ClickHouse Text-to-SQL Translator (Traducteur Métier).
Use Semantic Layer: {bool(config.get('use_semantic_layer'))}
Use Golden Queries: {bool(config.get('use_golden_queries'))}

CRITICAL TEXT-TO-SQL BEST PRACTICES:
- You will receive the DDL, column descriptions, and potentially "Golden Queries" as examples.
- Write precise ClickHouse SQL. Avoid complex joins if conditional aggregations (sumIf, countIf) suffice.
- If the query fails, you will receive the error and must auto-correct it (up to {config.get('max_retries') or 3} times).
- Generate a chart configuration (e.g., ECharts JSON) if requested and appropriate.

Return your response in JSON format: {{ "sql": "SELECT ...", "answer": "Explanation...", "chart_config": {{}} }}"""
        elif agent_type == "data_profiler_cleaner":
            system_prompt = f"""You are a Data Profiler & Cleaner (Data Steward).
Execution Engine: {config.get('execution_engine') or 'polars'}
Use Sampling: {bool(config.get('use_sampling'))} (Sample size: {config.get('sample_size') or 10000})

CRITICAL DATA PROFILING BEST PRACTICES:
- Review the provided statistical profile (nulls, cardinality, min/max, inferred types).
- Identify data quality issues (e.g., negative ages, inconsistent formats).
- Write a Python (Pandas/Polars) or SQL script to clean the data.
- NEVER execute destructive cleaning in production. Propose a script for human validation.

Return your response in JSON format: {{ "answer": "Audit summary...", "cleaning_script": "import polars as pl\\n..." }}"""
        elif agent_type == "unstructured_to_structured":
            schema = config.get("output_schema")
            schema_repr = schema if isinstance(schema, str) else json.dumps(schema or {})
            system_prompt = f"""You are a data extraction assistant. Convert the unstructured text into structured JSON.
Output Schema: {schema_repr}
Strict JSON: {config.get('strict_json') is not False}
Return ONLY valid JSON matching the schema."""
        elif agent_type == "email_cleaner":
            system_prompt = f"""You are an Email Cleaner. Condense the email into actionable sections.
Max bullets: {config.get('max_bullets') or 5}.
Include sections: {config.get('include_sections') or 'Action Items, Summary'}.
Return your response in JSON format: {{ "answer": "Cleaned email content..." }}"""
        elif agent_type == "file_assistant":
            system_prompt = f"You are a File Assistant. Answer questions based on the files in folder: {config.get('folder_path')}."
            folder_path = config.get("folder_path")
            if folder_path:
                try:
                    if os.path.exists(folder_path):
                        max_files = int(config.get("max_files") or 10)
                        files = os.listdir(folder_path)[:max_files]
                        extra_context = "\n\nFiles in directory:\n" + "\n".join(files)
                    else:
                        extra_context = f"\n\nError: Directory {folder_path} does not exist."
                except Exception as exc:
                    extra_context = f"\n\nError reading directory: {exc}"
        elif agent_type == "text_file_manager":
            text_result = await _execute_text_file_actions(user_input, config, on_step, agent_label)
            if text_result is not None:
                _emit_step(
                    on_step,
                    {
                        "status": "agent_completed",
                        "agent": agent_label,
                        "message": "Text file operations executed.",
                    },
                )
                return text_result

            system_prompt = (
                f"You are a Text File Manager. Folder: {config.get('folder_path')}. Allow overwrite: "
                f"{bool(config.get('allow_overwrite'))}.\n"
                "If the user asks to list/create/read/write/append/delete files, return JSON actions using this schema:\n"
                "{ \"actions\": [ {\"action\": \"list_files|create_file|read_file|write_file|append_file|delete_file\", "
                "\"file_path\": \"optional\", \"content\": \"optional\", \"max_chars\": 10000, \"recursive\": false, "
                "\"pattern\": \"optional\", \"prepend_newline\": false } ] }\n"
                "Otherwise return your response in JSON format: { \"answer\": \"File operations completed.\" }"
            )
        elif agent_type == "excel_manager":
            excel_result = await _execute_excel_actions(user_input, config, on_step, agent_label)
            if excel_result is not None:
                _emit_step(
                    on_step,
                    {
                        "status": "agent_completed",
                        "agent": agent_label,
                        "message": "Excel operations executed.",
                    },
                )
                return excel_result

            workbook_mode = "auto_in_folder" if _excel_manager_uses_auto_workbook_selection(config) else "fixed_path"
            sheet_mode = "auto_from_context" if _excel_manager_uses_auto_sheet_selection(config) else "fixed_default"
            workbook_label = str(config.get("workbook_path") or "").strip() or "auto"
            if workbook_mode == "auto_in_folder":
                workbook_label = "auto"
            sheet_label = str(config.get("default_sheet") or "").strip() or "Sheet1"
            if sheet_mode == "auto_from_context":
                sheet_label = "auto"
            system_prompt = (
                f"You are an Excel Manager. Folder: {config.get('folder_path')}. "
                f"Workbook preference: {workbook_label} ({workbook_mode}). "
                f"Default sheet preference: {sheet_label} ({sheet_mode}).\n"
                "If the user asks to create/modify/delete workbook/sheets/cells, return JSON actions using this schema:\n"
                "{ \"actions\": [ {\"action\": \"create_workbook|delete_workbook|list_sheets|create_sheet|delete_sheet|rename_sheet|write_cells|append_rows|read_sheet\", "
                "\"workbook_path\": \"optional\", \"sheet\": \"optional\", \"new_name\": \"optional\", \"cells\": {\"A1\": \"value\"}, "
                "\"rows\": [[\"v1\", \"v2\"]], \"max_rows\": 100 } ] }\n"
                "Otherwise return your response in JSON format: { \"answer\": \"Excel operations completed.\" }"
            )
        elif agent_type == "word_manager":
            word_result = await _execute_word_actions(user_input, config, on_step, agent_label)
            if word_result is not None:
                _emit_step(
                    on_step,
                    {
                        "status": "agent_completed",
                        "agent": agent_label,
                        "message": "Word operations executed.",
                    },
                )
                return word_result

            configured_document = str(config.get("document_path") or "").strip()
            auto_document_mode = _word_manager_uses_auto_document_selection(config)
            document_label = configured_document if configured_document else "auto"
            if auto_document_mode:
                document_label = "auto_in_folder"
            system_prompt = (
                f"You are a Word Manager. Folder: {config.get('folder_path')}. "
                f"Document preference: {document_label}.\n"
                "If the user asks to create/modify/delete/read document content, return JSON actions using this schema:\n"
                "{ \"actions\": [ {\"action\": \"create_document|delete_document|read_document|write_paragraphs|append_paragraphs|replace_text\", "
                "\"document_path\": \"optional\", \"paragraphs\": [\"line 1\", \"line 2\"], "
                "\"old_text\": \"optional\", \"new_text\": \"optional\", \"max_paragraphs\": 100 } ] }\n"
                "Otherwise return your response in JSON format: { \"answer\": \"Word operations completed.\" }"
            )
        elif agent_type == "elasticsearch_retriever":
            system_prompt = (
                f"You are an Elasticsearch Retriever. Index: {config.get('index')}. "
                "Synthesize the retrieved documents.\n"
                "Return your response in JSON format: { \"answer\": \"Synthesized content...\" }"
            )
            if config.get("base_url") and config.get("index"):
                try:
                    es_kwargs: Dict[str, Any] = {
                        "hosts": [config["base_url"]],
                        "verify_certs": config.get("verify_ssl") is not False,
                    }
                    if config.get("api_key"):
                        es_kwargs["api_key"] = config.get("api_key")
                    elif config.get("username"):
                        es_kwargs["basic_auth"] = (
                            config.get("username"),
                            config.get("password") or "",
                        )

                    client = Elasticsearch(**es_kwargs)
                    fields = ["*"]
                    if config.get("fields"):
                        fields = [field.strip() for field in str(config["fields"]).split(",") if field.strip()]

                    result = await asyncio.to_thread(
                        client.search,
                        index=config["index"],
                        query={
                            "simple_query_string": {
                                "query": user_input,
                                "fields": fields,
                            }
                        },
                        size=int(config.get("top_k") or 5),
                    )
                    hits = [hit.get("_source") for hit in result.get("hits", {}).get("hits", [])]
                    extra_context = "\n\nElasticsearch Results:\n" + json.dumps(hits)[:5000]
                except Exception as exc:
                    extra_context = f"\n\nElasticsearch Error: {exc}"
        elif agent_type == "rag_context":
            system_prompt = (
                f"You are a RAG Assistant. Folder: {config.get('folder_path')}. "
                f"Top K chunks: {config.get('top_k_chunks') or 3}.\n"
                "Return your response in JSON format: { \"answer\": \"RAG response...\" }"
            )
        elif agent_type == "rss_news":
            system_prompt = (
                f"You are an RSS News summarizer. Feeds: {config.get('feed_urls')}. "
                f"Interests: {config.get('interests')}."
            )
            configured_urls = [url.strip() for url in str(config.get("feed_urls") or "").split(",") if url.strip()]
            include_urls_from_question = bool(config.get("include_urls_from_question", True))
            question_urls = _extract_urls_from_text(user_input) if include_urls_from_question else []
            seed_urls: List[str] = []
            seen_seed: set[str] = set()
            for candidate in [*configured_urls, *question_urls]:
                normalized = candidate.rstrip("/")
                if normalized in seen_seed:
                    continue
                seen_seed.add(normalized)
                seed_urls.append(candidate)

            if seed_urls:
                try:
                    max_items_per_feed = max(1, min(int(config.get("max_items_per_feed") or 5), 30))
                    top_k = max(1, min(int(config.get("top_k") or 10), 100))
                    timeout_seconds = max(5, min(int(config.get("timeout_seconds") or 20), 45))
                    language_hint = _resolve_accept_language(config.get("language_hint"))
                    allow_reader_proxy = bool(config.get("allow_reader_proxy", True))
                    allow_google_news_fallback = bool(config.get("google_news_fallback", True))
                    interests = [
                        token.strip().lower()
                        for token in str(config.get("interests") or "").split(",")
                        if token.strip()
                    ]
                    exclude_keywords = [
                        token.strip().lower()
                        for token in str(config.get("exclude_keywords") or "").split(",")
                        if token.strip()
                    ]
                    include_general_if_no_match = bool(config.get("include_general_if_no_match", True))

                    harvested_items: List[Dict[str, str]] = []
                    failures: List[str] = []

                    for seed_url in seed_urls:
                        direct_response, direct_failures = await asyncio.to_thread(
                            _fetch_url_with_fallback,
                            seed_url,
                            timeout_seconds=timeout_seconds,
                            prefer_xml=True,
                            language_hint=language_hint,
                            allow_common_feed_candidates=True,
                            allow_google_news_fallback=allow_google_news_fallback,
                            allow_reader_proxy_fallback=allow_reader_proxy,
                        )
                        failures.extend(direct_failures[:2])

                        direct_entries: List[Dict[str, Any]] = []
                        direct_base_url = seed_url
                        discovered_feeds: List[str] = []
                        if direct_response is not None:
                            direct_base_url = direct_response.url or seed_url
                            direct_entries = _extract_feed_entries_from_content(direct_response.content)
                            if len(direct_entries) == 0:
                                discovered_feeds.extend(
                                    _discover_feed_links_from_html(
                                        direct_base_url,
                                        direct_response.text,
                                        max_links=8,
                                    )
                                )
                        if len(direct_entries) == 0 and len(discovered_feeds) == 0:
                            html_response, html_failures = await asyncio.to_thread(
                                _fetch_url_with_fallback,
                                seed_url,
                                timeout_seconds=timeout_seconds,
                                prefer_xml=False,
                                language_hint=language_hint,
                                allow_common_feed_candidates=False,
                                allow_google_news_fallback=False,
                                allow_reader_proxy_fallback=allow_reader_proxy,
                            )
                            failures.extend(html_failures[:1])
                            if html_response is not None:
                                discovered_feeds.extend(
                                    _discover_feed_links_from_html(
                                        html_response.url or seed_url,
                                        html_response.text,
                                        max_links=8,
                                    )
                                )

                        if len(direct_entries) > 0:
                            for entry in direct_entries[:max_items_per_feed]:
                                mapped = _entry_to_link_item(entry, direct_base_url)
                                if mapped is None:
                                    continue
                                summary = str(entry.get("summary") or entry.get("description") or "")
                                summary_text = BeautifulSoup(summary, "html.parser").get_text(" ", strip=True)
                                harvested_items.append(
                                    {
                                        "title": mapped["title"],
                                        "url": mapped["url"],
                                        "summary": summary_text,
                                        "source": seed_url,
                                    }
                                )
                            continue

                        if len(discovered_feeds) == 0:
                            discovered_feeds = _build_common_feed_candidates(seed_url, max_links=8)
                            if allow_google_news_fallback:
                                discovered_feeds.extend(
                                    _build_google_news_domain_feed_candidates(seed_url, max_links=4)
                                )

                        seen_discovered: set[str] = set()
                        for feed_url in discovered_feeds:
                            normalized_feed = str(feed_url or "").rstrip("/")
                            if not normalized_feed or normalized_feed in seen_discovered:
                                continue
                            seen_discovered.add(normalized_feed)
                            feed_response, feed_failures = await asyncio.to_thread(
                                _fetch_url_with_fallback,
                                feed_url,
                                timeout_seconds=timeout_seconds,
                                prefer_xml=True,
                                language_hint=language_hint,
                                allow_common_feed_candidates=False,
                                allow_google_news_fallback=False,
                                allow_reader_proxy_fallback=allow_reader_proxy,
                            )
                            failures.extend(feed_failures[:1])
                            if feed_response is None:
                                continue
                            feed_entries = _extract_feed_entries_from_content(feed_response.content)
                            if len(feed_entries) == 0:
                                continue
                            for entry in feed_entries[:max_items_per_feed]:
                                mapped = _entry_to_link_item(entry, feed_response.url or feed_url)
                                if mapped is None:
                                    continue
                                summary = str(entry.get("summary") or entry.get("description") or "")
                                summary_text = BeautifulSoup(summary, "html.parser").get_text(" ", strip=True)
                                harvested_items.append(
                                    {
                                        "title": mapped["title"],
                                        "url": mapped["url"],
                                        "summary": summary_text,
                                        "source": feed_response.url or feed_url,
                                    }
                                )
                            break

                    deduped_items: List[Dict[str, str]] = []
                    seen_item_keys: set[str] = set()
                    for item in harvested_items:
                        title = " ".join(str(item.get("title") or "").split()).strip()
                        url = str(item.get("url") or "").strip()
                        if not title or not url:
                            continue
                        dedupe_key = f"{title.lower()}|{url.rstrip('/').lower()}"
                        if dedupe_key in seen_item_keys:
                            continue
                        seen_item_keys.add(dedupe_key)
                        deduped_items.append(item)

                    matched_items: List[Dict[str, str]] = []
                    if len(interests) > 0:
                        for item in deduped_items:
                            searchable = (
                                f"{item.get('title', '')} {item.get('summary', '')}".lower()
                            )
                            if any(token in searchable for token in interests):
                                matched_items.append(item)
                    else:
                        matched_items = deduped_items

                    selected_items = matched_items
                    if len(selected_items) == 0 and include_general_if_no_match:
                        selected_items = deduped_items

                    filtered_items: List[Dict[str, str]] = []
                    for item in selected_items:
                        searchable = f"{item.get('title', '')} {item.get('summary', '')}".lower()
                        if any(token in searchable for token in exclude_keywords):
                            continue
                        filtered_items.append(item)

                    payload_items = filtered_items[:top_k]
                    if len(payload_items) > 0:
                        lines: List[str] = []
                        for item in payload_items:
                            title = item.get("title") or ""
                            summary = item.get("summary") or ""
                            url = item.get("url") or ""
                            source = item.get("source") or ""
                            if summary:
                                lines.append(f"- {title} ({url}) [{source}] :: {summary}")
                            else:
                                lines.append(f"- {title} ({url}) [{source}]")
                        extra_context = "\n\nRecent News Items:\n" + "\n".join(lines)
                    else:
                        details = " | ".join(failures[:3]) if failures else "No parsable RSS entries found."
                        return {
                            "answer": (
                                "RSS fetch failed for the configured sources. "
                                "This environment may be blocked by site protections (403), invalid feed URLs (404), "
                                f"or missing discoverable feed endpoints. Details: {details}"
                            ),
                            "details": {
                                "seed_urls": seed_urls,
                                "failures": failures[:10],
                            },
                        }
                except Exception as exc:
                    return {
                        "answer": f"RSS agent failed while fetching feeds: {exc}",
                        "details": {"seed_urls": seed_urls},
                    }
        elif agent_type == "web_scraper":
            system_prompt = (
                f"You are a Web Scraper. Start URLs: {config.get('start_urls')}. "
                f"Allowed domains: {config.get('allowed_domains')}."
            )
            start_urls = [url.strip() for url in str(config.get("start_urls") or "").split(",") if url.strip()]
            include_urls_from_question = bool(config.get("include_urls_from_question", True))
            question_urls = _extract_urls_from_text(user_input) if include_urls_from_question else []
            combined_urls: List[str] = []
            seen_combined: set[str] = set()
            for candidate in [*start_urls, *question_urls]:
                normalized = candidate.rstrip("/")
                if normalized in seen_combined:
                    continue
                seen_combined.add(normalized)
                combined_urls.append(candidate)

            if combined_urls:
                try:
                    max_chars = int(config.get("max_chars_per_page") or 5000)
                    timeout = int(config.get("timeout_seconds") or 15)
                    allow_reader_proxy = bool(config.get("allow_reader_proxy", True))
                    scraped = ""
                    failures: List[str] = []

                    for url in combined_urls:
                        try:
                            response, fetch_failures = await asyncio.to_thread(
                                _fetch_url_with_fallback,
                                url,
                                timeout_seconds=timeout,
                                prefer_xml=False,
                                language_hint=_resolve_accept_language(config.get("language_hint")),
                                allow_common_feed_candidates=False,
                                allow_google_news_fallback=False,
                                allow_reader_proxy_fallback=allow_reader_proxy,
                            )
                            failures.extend(fetch_failures[:1])
                            if response is None:
                                continue
                            soup = BeautifulSoup(response.text, "html.parser")
                            body_text = soup.get_text(" ", strip=True)
                            scraped += f"\n\nContent from {url}:\n{body_text[:max_chars]}"
                        except Exception:
                            continue

                    if scraped.strip():
                        extra_context = "\n\nScraped Content:\n" + scraped
                    elif failures:
                        return {
                            "answer": (
                                "Web scraper could not fetch content from configured URLs. "
                                f"Details: {' | '.join(failures[:3])}"
                            ),
                            "details": {"failures": failures[:10], "requested_urls": combined_urls},
                        }
                except Exception:
                    pass
        elif agent_type == "web_navigator":
            raw_tokens = _extract_raw_navigation_tokens(user_input)
            navigation_urls, invalid_navigation_urls = _normalize_navigation_url_candidates(raw_tokens)
            if not navigation_urls:
                start_url = str(config.get("start_url") or "").strip()
                if start_url:
                    fallback_urls, invalid_fallback = _normalize_navigation_url_candidates([start_url])
                    navigation_urls = fallback_urls
                    invalid_navigation_urls.extend(invalid_fallback)

            max_nav_steps = max(1, min(int(config.get("max_steps") or 5), 10))
            default_item_limit = int(config.get("max_links") or config.get("max_items") or 10)
            item_limit = _extract_requested_item_limit(user_input, default_limit=default_item_limit, hard_max=50)
            timeout_seconds = int(config.get("timeout_seconds") or 20)
            same_domain_only = config.get("same_domain_only")
            if same_domain_only is None:
                same_domain_only = True
            same_domain_only = bool(same_domain_only)
            allow_reader_proxy = bool(config.get("allow_reader_proxy", True))
            allow_google_news_fallback = bool(config.get("google_news_fallback", True))

            if invalid_navigation_urls:
                _emit_step(
                    on_step,
                    {
                        "status": "navigator_url_skipped",
                        "agent": agent_label,
                        "invalid_urls": invalid_navigation_urls[:10],
                        "message": "Ignored malformed URL candidates before navigation.",
                    },
                )

            if len(navigation_urls) == 0:
                details = (
                    f" Invalid URL(s): {', '.join(invalid_navigation_urls[:5])}."
                    if invalid_navigation_urls
                    else ""
                )
                _emit_step(
                    on_step,
                    {
                        "status": "agent_completed",
                        "agent": agent_label,
                        "message": "No URL provided to web navigator.",
                    },
                )
                return {
                    "answer": (
                        "No valid URL provided. Please include a full URL with host "
                        "(e.g. https://example.com) or configure `start_url`."
                        + details
                    )
                }

            collected_links: List[Dict[str, str]] = []
            seen_link_urls: set[str] = set()
            failures: List[str] = []
            requested_urls = navigation_urls[:max_nav_steps]
            language_hint = _resolve_accept_language(config.get("language_hint"))

            for target_url in requested_urls:
                _emit_step(
                    on_step,
                    {
                        "status": "navigator_fetch_started",
                        "agent": agent_label,
                        "url": target_url,
                    },
                )
                try:
                    response, fetch_failures = await asyncio.to_thread(
                        _fetch_url_with_fallback,
                        target_url,
                        timeout_seconds=timeout_seconds,
                        prefer_xml=False,
                        language_hint=language_hint,
                        allow_common_feed_candidates=False,
                        allow_google_news_fallback=False,
                        allow_reader_proxy_fallback=allow_reader_proxy,
                    )
                    failures.extend(fetch_failures[:2])
                    if response is None:
                        fallback_feed_candidates: List[str] = []
                        fallback_feed_candidates.extend(_build_common_feed_candidates(target_url, max_links=6))
                        if allow_google_news_fallback:
                            fallback_feed_candidates.extend(
                                _build_google_news_domain_feed_candidates(target_url, max_links=4)
                            )
                        for feed_url in fallback_feed_candidates:
                            if len(collected_links) >= item_limit:
                                break
                            feed_response, feed_failures = await asyncio.to_thread(
                                _fetch_url_with_fallback,
                                feed_url,
                                timeout_seconds=timeout_seconds,
                                prefer_xml=True,
                                language_hint=language_hint,
                                allow_common_feed_candidates=False,
                                allow_google_news_fallback=False,
                                allow_reader_proxy_fallback=allow_reader_proxy,
                            )
                            failures.extend(feed_failures[:1])
                            if feed_response is None:
                                continue
                            entries = _extract_feed_entries_from_content(feed_response.content)
                            if len(entries) == 0:
                                continue
                            remaining_from_feed = max(item_limit - len(collected_links), 0)
                            for entry in entries[:remaining_from_feed]:
                                mapped = _entry_to_link_item(entry, feed_response.url or feed_url)
                                if mapped is None:
                                    continue
                                mapped_url = mapped["url"]
                                parsed = urlparse(mapped_url)
                                if same_domain_only:
                                    base_domain = urlparse(target_url).netloc.lower()
                                    target_domain = parsed.netloc.lower()
                                    if (
                                        target_domain
                                        and base_domain
                                        and target_domain != base_domain
                                        and not target_domain.endswith(f".{base_domain}")
                                    ):
                                        continue
                                normalized_url = mapped_url.rstrip("/")
                                if normalized_url in seen_link_urls:
                                    continue
                                seen_link_urls.add(normalized_url)
                                collected_links.append(mapped)
                                if len(collected_links) >= item_limit:
                                    break
                            if len(collected_links) >= item_limit:
                                break

                        _emit_step(
                            on_step,
                            {
                                "status": "navigator_fetch_failed",
                                "agent": agent_label,
                                "url": target_url,
                                "error": fetch_failures[0] if fetch_failures else "Fetch failed",
                            },
                        )
                        continue

                    resolved_url = response.url or target_url
                    remaining = max(item_limit - len(collected_links), 0)
                    if remaining <= 0:
                        break

                    links = _extract_article_links_from_html(
                        resolved_url,
                        response.text,
                        max_links=remaining,
                        same_domain_only=same_domain_only,
                    )
                    if len(links) == 0:
                        links = _extract_fallback_links_from_html(
                            resolved_url,
                            response.text,
                            max_links=remaining,
                            same_domain_only=same_domain_only,
                        )

                    if len(links) == 0:
                        feed_candidates = _discover_feed_links_from_html(resolved_url, response.text, max_links=6)
                        if len(feed_candidates) == 0:
                            feed_candidates = _build_common_feed_candidates(resolved_url, max_links=6)
                            if allow_google_news_fallback:
                                feed_candidates.extend(
                                    _build_google_news_domain_feed_candidates(resolved_url, max_links=4)
                                )
                        for feed_url in feed_candidates:
                            if len(links) >= remaining:
                                break
                            feed_response, feed_failures = await asyncio.to_thread(
                                _fetch_url_with_fallback,
                                feed_url,
                                timeout_seconds=timeout_seconds,
                                prefer_xml=True,
                                language_hint=language_hint,
                                allow_common_feed_candidates=False,
                                allow_google_news_fallback=False,
                                allow_reader_proxy_fallback=allow_reader_proxy,
                            )
                            failures.extend(feed_failures[:1])
                            if feed_response is None:
                                continue
                            entries = _extract_feed_entries_from_content(feed_response.content)
                            if len(entries) == 0:
                                continue
                            for entry in entries:
                                mapped = _entry_to_link_item(entry, feed_response.url or feed_url)
                                if mapped is None:
                                    continue
                                mapped_url = mapped["url"]
                                parsed = urlparse(mapped_url)
                                if same_domain_only:
                                    base_domain = urlparse(resolved_url).netloc.lower()
                                    target_domain = parsed.netloc.lower()
                                    if (
                                        target_domain
                                        and base_domain
                                        and target_domain != base_domain
                                        and not target_domain.endswith(f".{base_domain}")
                                    ):
                                        continue
                                links.append(mapped)
                                if len(links) >= remaining:
                                    break
                    if len(links) == 0 and response.text:
                        links = _extract_links_from_text_blob(
                            resolved_url,
                            response.text,
                            max_links=remaining,
                            same_domain_only=same_domain_only,
                        )

                    pruned_links: List[Dict[str, str]] = []
                    for link in links:
                        normalized_url = str(link.get("url") or "").rstrip("/")
                        if not normalized_url or normalized_url in seen_link_urls:
                            continue
                        seen_link_urls.add(normalized_url)
                        pruned_links.append(link)
                        if len(pruned_links) >= remaining:
                            break
                    collected_links.extend(pruned_links)
                    _emit_step(
                        on_step,
                        {
                            "status": "navigator_fetch_completed",
                            "agent": agent_label,
                            "url": resolved_url,
                            "links_found": len(pruned_links),
                        },
                    )

                    if len(collected_links) >= item_limit:
                        break
                except Exception as exc:
                    failures.append(f"{target_url} -> {exc}")
                    _emit_step(
                        on_step,
                        {
                            "status": "navigator_fetch_failed",
                            "agent": agent_label,
                            "url": target_url,
                            "error": str(exc),
                        },
                    )

            if len(collected_links) > 0:
                payload = collected_links[:item_limit]
                answer = json.dumps(payload, ensure_ascii=False, indent=2)
                _emit_step(
                    on_step,
                    {
                        "status": "agent_completed",
                        "agent": agent_label,
                        "message": f"Navigation completed with {len(payload)} links.",
                    },
                )
                return {
                    "answer": answer,
                    "details": {
                        "links": payload,
                        "requested_urls": requested_urls,
                        "invalid_urls": invalid_navigation_urls,
                        "failures": failures,
                    },
                }

            fallback_message = (
                "Web navigation failed to extract article links from the target page(s). "
                "The site may block automated requests or require JavaScript rendering."
            )
            if failures:
                fallback_message += f" Details: {' | '.join(failures[:3])}"

            _emit_step(
                on_step,
                {
                    "status": "agent_completed",
                    "agent": agent_label,
                    "message": "Navigation finished with no extractable links.",
                },
            )
            return {
                "answer": fallback_message,
                "details": {
                    "requested_urls": requested_urls,
                    "invalid_urls": invalid_navigation_urls,
                    "failures": failures,
                },
            }
        else:
            system_prompt = f"""Role: {config.get('role') or 'Assistant'}
Persona: {config.get('persona') or 'Helpful'}
Objectives: {config.get('objectives') or 'Assist the user'}
Instructions: {config.get('system_prompt') or ''}"""

        try:
            _emit_step(
                on_step,
                {
                    "status": "llm_call_started",
                    "agent": agent_label,
                    "message": "Generating response with the language model.",
                },
            )
            response = await llm.ainvoke(
                [
                    SystemMessage(content=system_prompt),
                    HumanMessage(content=user_input + extra_context),
                ]
            )
            content = _coerce_message_content(response.content)

            _emit_step(
                on_step,
                {
                    "status": "llm_response_received",
                    "agent": agent_label,
                    "message": "Model response received. Preparing final answer.",
                },
            )

            if agent_type == "excel_manager":
                excel_llm_result = await _execute_excel_actions(
                    content,
                    config,
                    on_step,
                    agent_label,
                    allow_heuristics=False,
                )
                if excel_llm_result is not None:
                    _emit_step(
                        on_step,
                        {
                            "status": "agent_completed",
                            "agent": agent_label,
                            "message": "Excel operations executed.",
                        },
                    )
                    return excel_llm_result
            if agent_type == "text_file_manager":
                text_llm_result = await _execute_text_file_actions(
                    content,
                    config,
                    on_step,
                    agent_label,
                    allow_heuristics=False,
                )
                if text_llm_result is not None:
                    _emit_step(
                        on_step,
                        {
                            "status": "agent_completed",
                            "agent": agent_label,
                            "message": "Text file operations executed.",
                        },
                    )
                    return text_llm_result
            if agent_type == "word_manager":
                word_llm_result = await _execute_word_actions(
                    content,
                    config,
                    on_step,
                    agent_label,
                    allow_heuristics=False,
                )
                if word_llm_result is not None:
                    _emit_step(
                        on_step,
                        {
                            "status": "agent_completed",
                            "agent": agent_label,
                            "message": "Word operations executed.",
                        },
                    )
                    return word_llm_result

            if content.strip().startswith("{") or agent_type != "custom":
                try:
                    json_payload = json.loads(_extract_json_block(content))
                    result: AgentResult = {
                        "answer": json_payload.get("answer")
                        or (json.dumps(json_payload) if agent_type == "unstructured_to_structured" else "JSON parsed successfully."),
                        "sql": json_payload.get("sql"),
                        "rows": json_payload.get("rows") or [],
                        "details": json_payload,
                    }
                    _emit_step(
                        on_step,
                        {
                            "status": "agent_completed",
                            "agent": agent_label,
                            "message": "Response ready.",
                        },
                    )
                    return result
                except Exception:
                    if agent_type == "unstructured_to_structured":
                        fallback_result = {"answer": f"Fallback raw output:\n{content}"}
                        _emit_step(
                            on_step,
                            {
                                "status": "agent_completed",
                                "agent": agent_label,
                                "message": "Response ready.",
                            },
                        )
                        return fallback_result
                    plain_result = {"answer": content}
                    _emit_step(
                        on_step,
                        {
                            "status": "agent_completed",
                            "agent": agent_label,
                            "message": "Response ready.",
                        },
                    )
                    return plain_result

            plain_result = {"answer": content}
            _emit_step(
                on_step,
                {
                    "status": "agent_completed",
                    "agent": agent_label,
                    "message": "Response ready.",
                },
            )
            return plain_result
        except Exception as exc:
            _emit_step(
                on_step,
                {
                    "status": "agent_error",
                    "agent": agent_label,
                    "error": str(exc),
                    "message": "Agent execution failed.",
                },
            )
            return {"answer": f"Error executing agent: {exc}"}


def _runtime_unavailable_reason_for_agent(agent: Dict[str, Any]) -> str | None:
    agent_type = str(agent.get("agent_type") or "").strip().lower()
    config = build_agent_execution_config(agent)

    if agent_type == "clickhouse_specific":
        templates = _normalize_clickhouse_specific_templates(config)
        if len(templates) == 0:
            return "Missing ClickHouse templates (`query_templates`) configuration."

    if agent_type == "excel_manager":
        try:
            import openpyxl  # type: ignore  # noqa: F401
        except Exception:
            return "Dependency missing: openpyxl is required."

    if agent_type == "word_manager":
        try:
            from docx import Document  # type: ignore  # noqa: F401
        except Exception:
            return "Dependency missing: python-docx is required."

    if agent_type == "elasticsearch_retriever":
        if not str(config.get("base_url") or "").strip() or not str(config.get("index") or "").strip():
            return "Missing Elasticsearch configuration (base_url/index)."

    return None


def _collect_runtime_unavailable_agents(agents: List[Dict[str, Any]]) -> Dict[int, str]:
    unavailable: Dict[int, str] = {}
    for agent in agents:
        reason = _runtime_unavailable_reason_for_agent(agent)
        if not reason:
            continue
        agent_id = _safe_int(agent.get("id"))
        if agent_id is None:
            continue
        unavailable[agent_id] = reason
    return unavailable


def _normalize_manager_call_input(value: str) -> str:
    normalized = re.sub(r"\s+", " ", str(value or "").strip()).lower()
    return normalized[:600]


def _extract_manager_call_context(call: Dict[str, Any]) -> Dict[str, Any]:
    context: Dict[str, Any] = {}
    query_name = str(call.get("query_name") or call.get("template_name") or "").strip()
    if query_name:
        context["query_name"] = query_name

    params: Dict[str, Any] = {}
    params.update(_normalize_clickhouse_specific_params(call.get("params")))
    params.update(_normalize_clickhouse_specific_params(call.get("parameters")))
    if len(params) > 0:
        context["params"] = params

    return context


def _manager_call_signature(
    agent_id: int | None,
    call_input: str,
    call_context: Dict[str, Any] | None = None,
) -> str:
    normalized_context = ""
    if isinstance(call_context, dict) and len(call_context) > 0:
        try:
            normalized_context = json.dumps(call_context, sort_keys=True, ensure_ascii=False)
        except Exception:
            normalized_context = str(call_context)
    return f"{agent_id or 0}|{_normalize_manager_call_input(call_input)}|{normalized_context}"


def _should_mark_agent_unavailable(error_text: str) -> bool:
    normalized = str(error_text or "").lower()
    permanent_markers = (
        "requires the `openpyxl` dependency",
        "requires the `python-docx` dependency",
        "no module named",
        "dependency missing",
        "unsupported llm provider",
    )
    return any(marker in normalized for marker in permanent_markers)


def _extract_terminal_failure_reason(text: str) -> str | None:
    normalized = str(text or "").strip().lower()
    if not normalized:
        return None

    if "workbook not found" in normalized and "auto_create_workbook" in normalized:
        return "excel_workbook_missing"
    if "requires the `openpyxl` dependency" in normalized or "requires the `python-docx` dependency" in normalized:
        return "dependency_missing"
    if "i’m sorry, but i can’t help with that" in normalized or "i'm sorry, but i can't help with that" in normalized:
        return "capability_refusal"
    if "can't browse the web" in normalized or "cannot browse the web" in normalized:
        return "external_access_blocked"
    if "cannot access external urls" in normalized or "cannot access external url" in normalized:
        return "external_access_blocked"
    if "http 403" in normalized or "403 forbidden" in normalized:
        return "http_forbidden"
    if "http 404" in normalized:
        return "http_not_found"
    if "invalid_url" in normalized:
        return "invalid_input_url"
    if "no valid url provided" in normalized:
        return "invalid_input_url"
    if "web navigation failed" in normalized and ("block" in normalized or "javascript rendering" in normalized):
        return "external_access_blocked"
    return None


def _terminal_failure_threshold(reason: str) -> int:
    thresholds = {
        "dependency_missing": 1,
        "capability_refusal": 1,
        "external_access_blocked": 2,
        "http_forbidden": 2,
        "http_not_found": 2,
        "invalid_input_url": 2,
        "excel_workbook_missing": 2,
    }
    return thresholds.get(reason, 2)


def _summarize_terminal_failures_for_user(failures: List[Dict[str, Any]]) -> str:
    if not failures:
        return "Task stopped because required agent actions repeatedly failed."

    lines = ["I could not finish automatically because required agent actions repeatedly failed:"]
    for failure in failures[:4]:
        agent = str(failure.get("agent") or "agent")
        reason = str(failure.get("reason") or "unknown_failure").replace("_", " ")
        detail = str(failure.get("detail") or "").strip()
        if detail:
            lines.append(f"- {agent}: {reason} ({detail})")
        else:
            lines.append(f"- {agent}: {reason}")

    lines.append("Please adjust configuration/inputs and retry.")
    return "\n".join(lines)


def _safe_json_loads(text: str) -> Dict[str, Any]:
    try:
        parsed = json.loads(_extract_json_block(text))
        if isinstance(parsed, dict):
            return parsed
    except Exception:
        pass
    return {}


def _normalize_plan_graph(raw_plan: Any) -> List[Dict[str, Any]]:
    if not isinstance(raw_plan, list):
        return []

    nodes: List[Dict[str, Any]] = []
    seen_ids: set[str] = set()
    for idx, item in enumerate(raw_plan, start=1):
        if not isinstance(item, dict):
            continue
        task_id = str(item.get("id") or item.get("task_id") or f"T{idx}").strip() or f"T{idx}"
        if task_id in seen_ids:
            continue
        seen_ids.add(task_id)
        description = str(item.get("task") or item.get("description") or "").strip()
        if not description:
            continue
        status = str(item.get("status") or "pending").strip().lower()
        if status not in {"pending", "running", "done", "blocked", "failed"}:
            status = "pending"
        depends_on_raw = item.get("depends_on") or item.get("dependencies") or []
        depends_on: List[str] = []
        if isinstance(depends_on_raw, list):
            for dependency in depends_on_raw:
                dep = str(dependency or "").strip()
                if dep and dep != task_id:
                    depends_on.append(dep)
        success_criteria = str(item.get("success_criteria") or item.get("done_when") or "").strip()
        nodes.append(
            {
                "id": task_id,
                "task": description,
                "depends_on": list(dict.fromkeys(depends_on)),
                "status": status,
                "success_criteria": success_criteria,
            }
        )
    return nodes


def _plan_graph_to_text(plan_graph: List[Dict[str, Any]]) -> str:
    if not plan_graph:
        return "No explicit DAG plan yet."
    lines: List[str] = []
    for node in plan_graph:
        depends = ", ".join(node.get("depends_on") or []) or "none"
        success = str(node.get("success_criteria") or "").strip()
        base = (
            f"- {node.get('id')}: {node.get('task')} "
            f"[status={node.get('status')}, depends_on={depends}]"
        )
        if success:
            base += f" | done_when={success}"
        lines.append(base)
    return "\n".join(lines)


def _update_history_summary(current_summary: str, new_entry: str, max_chars: int = 3200) -> str:
    base = str(current_summary or "").strip()
    append = str(new_entry or "").strip()
    if not append:
        return base
    merged = f"{base}\n{append}".strip() if base else append
    return _clip_text(merged, max_chars=max_chars)


def _default_definition_of_done(user_input: str) -> List[str]:
    return [
        "All required sub-goals from the user request are addressed.",
        "No blocking dependency/error remains unresolved without explanation.",
        "Final answer contains concrete outputs (results/files/links) or explicit missing-data reasons.",
        f"Task remains aligned with user objective: { _clip_text(user_input, 240) }",
    ]


def _apply_plan_progress_from_outcomes(
    plan_graph: List[Dict[str, Any]],
    outcomes: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    if not plan_graph:
        return []
    if not outcomes:
        return plan_graph

    updated = [dict(node) for node in plan_graph]
    if any(bool(item.get("informative")) for item in outcomes):
        for node in updated:
            if node.get("status") in {"pending", "running"}:
                node["status"] = "done"
                break
    elif all(item.get("terminal_reason") for item in outcomes):
        for node in updated:
            if node.get("status") in {"pending", "running"}:
                node["status"] = "blocked"
                break
    return updated


class ManagerState(TypedDict):
    input: str
    active_agents: List[Dict[str, Any]]
    llm: Any
    manager_config: Dict[str, Any]
    steps: List[Dict[str, Any]]
    conversation_history: str
    scratchpad: Dict[str, Any]
    current_step: int
    agent_calls_count: int
    max_agent_calls: int
    max_steps: int
    done: bool
    final_answer: str
    unavailable_agents: Dict[int, str]
    successful_call_signatures: set[str]
    attempted_call_signatures: set[str]
    current_plan: str
    plan_history: List[str]
    terminal_failure_counts: Dict[int, Dict[str, int]]
    history_summary: str
    flight_plan: List[Dict[str, Any]]
    plan_created: bool
    definition_of_done: List[str]


class MultiAgentManager:
    @staticmethod
    async def run_stream(
        user_input: str,
        agents: List[Dict[str, Any]],
        llm: Any,
        manager_config: Dict[str, Any] | None = None,
        on_step: Callable[[Dict[str, Any]], None] | None = None,
    ) -> Dict[str, Any]:
        cfg = manager_config or {}
        active_agents = [agent for agent in agents if agent.get("agent_type") != "manager"]
        unavailable_agents = _collect_runtime_unavailable_agents(active_agents)
        plannable_agents = [
            agent for agent in active_agents if _safe_int(agent.get("id")) not in unavailable_agents
        ]

        state: ManagerState = {
            "input": user_input,
            "active_agents": active_agents,
            "llm": llm,
            "manager_config": cfg,
            "steps": [],
            "conversation_history": f"User Request: {user_input}\n",
            "scratchpad": {},
            "current_step": 0,
            "agent_calls_count": 0,
            "max_agent_calls": max(1, int(cfg.get("max_agent_calls") or 15)),
            "max_steps": max(1, int(cfg.get("max_steps") or 8)),
            "done": False,
            "final_answer": "",
            "unavailable_agents": unavailable_agents,
            "successful_call_signatures": set(),
            "attempted_call_signatures": set(),
            "current_plan": "",
            "plan_history": [],
            "terminal_failure_counts": {},
            "history_summary": "",
            "flight_plan": [],
            "plan_created": False,
            "definition_of_done": _default_definition_of_done(user_input),
        }

        def add_step(step: Dict[str, Any]) -> None:
            state["steps"].append(step)
            if on_step:
                on_step(step)

        add_step(
            {
                "status": "manager_start",
                "message": "Starting multi-agent orchestration",
                "active_agents": len(active_agents),
                "plannable_agents": len(plannable_agents),
                "unavailable_agents": [
                    {
                        "id": agent.get("id"),
                        "name": agent.get("name"),
                        "agent_type": agent.get("agent_type"),
                        "reason": unavailable_agents.get(_safe_int(agent.get("id")) or -1),
                    }
                    for agent in active_agents
                    if (_safe_int(agent.get("id")) or -1) in unavailable_agents
                ],
                "max_steps": state["max_steps"],
                "max_agent_calls": state["max_agent_calls"],
                "definition_of_done": state["definition_of_done"],
            }
        )

        async def iterate(current_state: ManagerState) -> ManagerState:
            if current_state["done"]:
                return current_state

            if len(current_state["active_agents"]) == 0:
                current_state["final_answer"] = "No specialized agents are configured. Create at least one non-manager agent."
                add_step(
                    {
                        "status": "manager_final",
                        "answer": current_state["final_answer"],
                        "manager_summary": "No available worker agents.",
                    }
                )
                current_state["done"] = True
                return current_state

            current_plannable_agents = [
                agent
                for agent in current_state["active_agents"]
                if _safe_int(agent.get("id")) not in current_state["unavailable_agents"]
            ]
            if len(current_plannable_agents) == 0:
                current_state["final_answer"] = (
                    "No runnable agents are currently available. "
                    "Fix dependencies/configuration and retry."
                )
                add_step(
                    {
                        "status": "manager_final",
                        "answer": current_state["final_answer"],
                        "manager_summary": "All worker agents are unavailable at runtime.",
                        "unavailable_agents": [
                            {
                                "id": agent.get("id"),
                                "name": agent.get("name"),
                                "agent_type": agent.get("agent_type"),
                                "reason": current_state["unavailable_agents"].get(_safe_int(agent.get("id")) or -1),
                            }
                            for agent in current_state["active_agents"]
                            if (_safe_int(agent.get("id")) or -1) in current_state["unavailable_agents"]
                        ],
                    }
                )
                current_state["done"] = True
                return current_state

            if current_state["current_step"] >= current_state["max_steps"]:
                current_state["final_answer"] = "Reached maximum steps without a final answer."
                add_step(
                    {
                        "status": "manager_final",
                        "answer": current_state["final_answer"],
                        "manager_summary": "Max steps reached.",
                    }
                )
                current_state["done"] = True
                return current_state

            agent_catalog = [_agent_catalog_entry(agent) for agent in current_plannable_agents]
            unavailable_summary = [
                {
                    "id": agent.get("id"),
                    "name": agent.get("name"),
                    "agent_type": agent.get("agent_type"),
                    "reason": current_state["unavailable_agents"].get(_safe_int(agent.get("id")) or -1),
                }
                for agent in current_state["active_agents"]
                if (_safe_int(agent.get("id")) or -1) in current_state["unavailable_agents"]
            ]
            remaining_steps = max(current_state["max_steps"] - current_state["current_step"], 0)
            remaining_calls = max(current_state["max_agent_calls"] - current_state["agent_calls_count"], 0)
            previous_plan = _clip_text(str(current_state.get("current_plan") or ""), 2000) or "No prior plan yet."
            recent_trace = _build_recent_manager_trace(current_state.get("steps") or [])
            manager_policy_prompt = _clip_text(
                str(current_state.get("manager_config", {}).get("system_prompt") or ""),
                2400,
            )
            flight_plan_text = _plan_graph_to_text(current_state.get("flight_plan") or [])
            history_summary = (
                _clip_text(str(current_state.get("history_summary") or ""), 2600)
                or "No summarized history yet."
            )
            dod_text = "\n".join(f"- {item}" for item in (current_state.get("definition_of_done") or []))

            add_step(
                {
                    "status": "manager_iteration_start",
                    "iteration": current_state["current_step"] + 1,
                    "remaining_steps": remaining_steps,
                    "remaining_agent_calls": remaining_calls,
                    "current_plan": previous_plan,
                }
            )

            if remaining_steps <= 2 or remaining_calls <= 2:
                add_step(
                    {
                        "status": "manager_guardrail_warning",
                        "remaining_steps": remaining_steps,
                        "remaining_agent_calls": remaining_calls,
                        "message": (
                            "Budget is nearly exhausted. Finalize decisively and avoid optional calls."
                        ),
                    }
                )

            planner_prompt = f"""You are an advanced Multi-Agent Orchestrator using strict ReAct.

MANDATORY LOOP EACH ITERATION:
1) THOUGHT (private): reason about the objective and constraints.
2) ACTION: choose exactly one tool among create_plan, update_plan, dispatch_agents, finish_task.
3) OBSERVATION EXPECTED: explain what evidence should be produced by that action.

RULES:
- First iteration MUST call create_plan.
- Plan must be represented as DAG JSON (tasks + dependencies + statuses).
- Keep a compact scratchpad, not full transcript.
- Never dispatch an identical call already attempted.
- After each subordinate-agent result, re-evaluate whether the current step objective is met.
- If a tool fails, analyze cause and pick an alternative approach (no blind retry loops).
- finish_task is the only valid way to stop.
- Respect hard limits. If budget is low, prioritize closure.

Definition of Done (all criteria should be checked before finish_task):
{dod_text}

Available agents catalog:
{json.dumps(agent_catalog, ensure_ascii=False, indent=2)}

Unavailable agents (DO NOT CALL):
{json.dumps(unavailable_summary, ensure_ascii=False, indent=2)}

Manager operating prompt:
{manager_policy_prompt or "No custom manager prompt provided. Use best orchestration practices."}

Execution budget:
- max_steps: {current_state["max_steps"]}
- current_step: {current_state["current_step"]}
- remaining_steps: {remaining_steps}
- max_agent_calls: {current_state["max_agent_calls"]}
- agent_calls_used: {current_state["agent_calls_count"]}
- agent_calls_remaining: {remaining_calls}

Current flight plan (DAG):
{flight_plan_text}

Previous plan:
{previous_plan}

Recent trace:
{recent_trace}

Summarized history:
{history_summary}

Shared scratchpad:
{json.dumps(current_state['scratchpad'], indent=2)}

Return ONLY JSON:
{{
  "tool": "create_plan|update_plan|dispatch_agents|finish_task",
  "private_reasoning": "internal CoT, concise but explicit (not for end user display)",
  "observation_expected": "what evidence is expected after action",
  "current_plan": "updated flight plan summary",
  "dag_plan": [{{"id":"T1","task":"...","depends_on":[],"status":"pending","success_criteria":"..."}}],
  "plan_revision": "what changed vs previous plan",
  "focus_check": "one-line focus reminder",
  "rationale": "brief external-safe rationale",
  "scratchpad_updates": {{"key":"value"}},
  "calls": [{{"agent_id":1,"agent_name":"optional","agent_type":"optional","input":"targeted instruction","query_name":"optional","params":{{"P1":"value"}}}}],
  "definition_of_done_check": [{{"criterion":"...","satisfied":true,"evidence":"..."}}],
  "final_answer": "required when tool=finish_task",
  "missing_information": "optional"
}}"""

            try:
                planner_response = await current_state["llm"].ainvoke([HumanMessage(content=planner_prompt)])
                planner_content = _coerce_message_content(planner_response.content)
                planner_decision = _safe_json_loads(planner_content)

                if not planner_decision:
                    raise ValueError("Planner returned non-JSON or invalid decision.")

                selected_tool = str(planner_decision.get("tool") or "").strip().lower()
                if not current_state.get("plan_created") and selected_tool != "create_plan":
                    planner_decision["tool"] = "create_plan"
                    if not isinstance(planner_decision.get("dag_plan"), list):
                        planner_decision["dag_plan"] = [
                            {
                                "id": "T1",
                                "task": "Understand the request and create an execution plan.",
                                "depends_on": [],
                                "status": "pending",
                                "success_criteria": "A valid DAG plan exists.",
                            },
                            {
                                "id": "T2",
                                "task": "Execute the plan with relevant agents.",
                                "depends_on": ["T1"],
                                "status": "pending",
                                "success_criteria": "Evidence gathered from agent outputs.",
                            },
                            {
                                "id": "T3",
                                "task": "Synthesize and finalize answer with finish_task.",
                                "depends_on": ["T2"],
                                "status": "pending",
                                "success_criteria": "Final answer submitted.",
                            },
                        ]
                    planner_decision["plan_revision"] = "create_plan enforced because plan creation is mandatory first."
                    selected_tool = "create_plan"

                critic_prompt = f"""You are a strict execution critic for a multi-agent manager.
Review the planner decision and return whether it should be approved or revised.

Rules:
- Enforce ReAct consistency (thought -> action -> expected observation).
- Ensure first action is create_plan if no flight plan exists.
- Ensure chosen action respects remaining step/call budget.
- Prevent redundant calls or blind retries.
- If planner asks to finish, verify Definition of Done evidence is sufficient.

Current context summary:
- remaining_steps: {remaining_steps}
- remaining_agent_calls: {remaining_calls}
- plan_created: {bool(current_state.get("plan_created"))}
- previous_plan: {previous_plan}
- flight_plan: {flight_plan_text}
- history_summary: {history_summary}

Planner decision JSON:
{json.dumps(planner_decision, ensure_ascii=False, indent=2)}

Return ONLY JSON:
{{
  "verdict": "approve|revise",
  "feedback": "short critical feedback",
  "revised_decision": {{
    "tool": "create_plan|update_plan|dispatch_agents|finish_task",
    "current_plan": "optional",
    "dag_plan": [{{"id":"T1","task":"...","depends_on":[],"status":"pending","success_criteria":"..."}}],
    "calls": [],
    "final_answer": "",
    "focus_check": "",
    "rationale": "",
    "plan_revision": "",
    "definition_of_done_check": []
  }}
}}"""
                critic_response = await current_state["llm"].ainvoke([HumanMessage(content=critic_prompt)])
                critic_content = _coerce_message_content(critic_response.content)
                critic_decision = _safe_json_loads(critic_content)

                decision: Dict[str, Any] = dict(planner_decision)
                critic_verdict = str(critic_decision.get("verdict") or "").strip().lower()
                revised_decision = critic_decision.get("revised_decision")
                if critic_verdict == "revise" and isinstance(revised_decision, dict):
                    decision = {**decision, **revised_decision}
                    add_step(
                        {
                            "status": "manager_critique",
                            "verdict": "revise",
                            "feedback": _clip_text(str(critic_decision.get("feedback") or ""), 800),
                        }
                    )
                else:
                    add_step(
                        {
                            "status": "manager_critique",
                            "verdict": "approve",
                            "feedback": _clip_text(str(critic_decision.get("feedback") or ""), 800),
                        }
                    )

                tool = str(decision.get("tool") or "").strip().lower()
                if tool in {"dispatch_agents", "dispatch", "call_agents", "act"}:
                    decision["status"] = "calling_agent"
                elif tool in {"finish_task", "submit_final_answer"}:
                    decision["status"] = "final_answer"
                else:
                    decision["status"] = "thinking"

                normalized_plan_graph = _normalize_plan_graph(decision.get("dag_plan"))
                if tool in {"create_plan", "update_plan"} and normalized_plan_graph:
                    current_state["flight_plan"] = normalized_plan_graph
                    current_state["plan_created"] = True
                elif tool == "create_plan" and not normalized_plan_graph and not current_state["plan_created"]:
                    fallback_plan = _normalize_plan_graph(
                        [
                            {
                                "id": "T1",
                                "task": "Gather required evidence via specialized agents.",
                                "depends_on": [],
                                "status": "pending",
                                "success_criteria": "Evidence collected.",
                            },
                            {
                                "id": "T2",
                                "task": "Synthesize and finish task.",
                                "depends_on": ["T1"],
                                "status": "pending",
                                "success_criteria": "Final answer submitted with evidence.",
                            },
                        ]
                    )
                    current_state["flight_plan"] = fallback_plan
                    current_state["plan_created"] = True

                updates = decision.get("scratchpad_updates")
                if isinstance(updates, dict):
                    current_state["scratchpad"] = {**current_state["scratchpad"], **updates}

                rationale = _clip_text(str(decision.get("rationale") or ""), 2000)
                raw_plan = _clip_text(str(decision.get("current_plan") or ""), 3000)
                if not raw_plan and current_state.get("flight_plan"):
                    raw_plan = _plan_graph_to_text(current_state["flight_plan"])
                previous_plan = _clip_text(str(current_state.get("current_plan") or ""), 3000)
                plan = raw_plan or previous_plan or "No explicit plan provided."
                plan_revision = _clip_text(str(decision.get("plan_revision") or ""), 1200) or _plan_revision_summary(
                    previous_plan,
                    plan,
                )
                focus_check = _clip_text(str(decision.get("focus_check") or ""), 1200)
                current_state["current_plan"] = plan
                current_state["plan_history"].append(plan)

                dod_checks = decision.get("definition_of_done_check")
                if (
                    decision.get("status") == "final_answer"
                    and isinstance(dod_checks, list)
                    and remaining_steps > 1
                ):
                    unsatisfied = [
                        item
                        for item in dod_checks
                        if isinstance(item, dict) and not bool(item.get("satisfied"))
                    ]
                    if unsatisfied:
                        decision["status"] = "thinking"
                        decision["tool"] = "update_plan"
                        add_step(
                            {
                                "status": "manager_warning",
                                "message": (
                                    "finish_task rejected because Definition of Done is not fully satisfied yet."
                                ),
                                "unsatisfied_criteria": unsatisfied[:5],
                            }
                        )

                add_step(
                    {
                        "status": "manager_decision",
                        "tool": str(decision.get("tool") or ""),
                        "observation_expected": _clip_text(str(decision.get("observation_expected") or ""), 1200),
                        "rationale": rationale,
                        "plan": plan,
                        "plan_revision": plan_revision,
                        "focus_check": focus_check,
                    }
                )
                current_state["conversation_history"] += (
                    f"\n[Manager Plan Updated]: {plan}\n"
                    f"[Manager Plan Revision]: {plan_revision}\n"
                    f"[Manager Focus Check]: {focus_check}\n"
                    f"[Manager Rationale]: {rationale}\n"
                )
                current_state["history_summary"] = _update_history_summary(
                    current_state.get("history_summary") or "",
                    (
                        f"Step {current_state['current_step'] + 1}: "
                        f"tool={decision.get('tool')} | plan_revision={plan_revision} | focus={focus_check}"
                    ),
                )

                if decision.get("status") == "final_answer":
                    current_state["final_answer"] = decision.get("final_answer") or ""
                    add_step(
                        {
                            "status": "manager_final",
                            "answer": current_state["final_answer"],
                            "manager_summary": rationale,
                            "judge_verdict": "Pass",
                            "missing_information": decision.get("missing_information"),
                        }
                    )
                    current_state["done"] = True
                elif (
                    decision.get("status") == "calling_agent"
                    and isinstance(decision.get("calls"), list)
                    and len(decision["calls"]) > 0
                ):
                    remaining_agent_calls = current_state["max_agent_calls"] - current_state["agent_calls_count"]
                    if remaining_agent_calls <= 0:
                        current_state["final_answer"] = "Reached max agent-call budget. Provide a final answer from gathered evidence."
                        add_step(
                            {
                                "status": "manager_final",
                                "answer": current_state["final_answer"],
                                "manager_summary": "Max agent-call budget reached.",
                            }
                        )
                        current_state["done"] = True
                    else:
                        raw_calls = [call for call in decision["calls"] if isinstance(call, dict)]
                        bounded_calls = raw_calls[:remaining_agent_calls]

                        unique_calls: List[Dict[str, Any]] = []
                        seen_signatures: set[str] = set()
                        for call in bounded_calls:
                            call_context = _extract_manager_call_context(call)
                            context_signature = ""
                            if len(call_context) > 0:
                                try:
                                    context_signature = json.dumps(call_context, sort_keys=True, ensure_ascii=False)
                                except Exception:
                                    context_signature = str(call_context)
                            signature = "|".join(
                                [
                                    str(call.get("agent_id") or ""),
                                    _normalize_identifier(call.get("agent_name") or call.get("agent") or call.get("agent_type")),
                                    str(call.get("input") or "").strip(),
                                    context_signature,
                                ]
                            )
                            if signature in seen_signatures:
                                continue
                            seen_signatures.add(signature)
                            unique_calls.append(call)

                        resolved_calls: List[tuple[Dict[str, Any], Dict[str, Any], str, Dict[str, Any], str]] = []
                        for call in unique_calls:
                            selected = _resolve_agent_from_call(call, current_plannable_agents)
                            if selected is None:
                                add_step(
                                    {
                                        "status": "agent_not_found",
                                        "requested_call": call,
                                        "available_agents": [
                                            {
                                                "id": agent.get("id"),
                                                "name": agent.get("name"),
                                                "agent_type": agent.get("agent_type"),
                                            }
                                            for agent in current_plannable_agents
                                        ],
                                    }
                                )
                                continue
                            call_input = str(call.get("input") or "").strip() or str(current_state["input"])
                            call_context = _extract_manager_call_context(call)
                            selected_id = _safe_int(selected.get("id"))
                            signature = _manager_call_signature(selected_id, call_input, call_context)
                            if signature in current_state["attempted_call_signatures"]:
                                add_step(
                                    {
                                        "status": "manager_warning",
                                        "message": (
                                            "Skipping duplicate call because an equivalent call context "
                                            f"was already attempted for agent '{selected.get('name')}'."
                                        ),
                                        "agent": selected.get("name"),
                                        "agent_id": selected_id,
                                        "input": call_input,
                                        "query_name": call_context.get("query_name"),
                                        "params": call_context.get("params"),
                                    }
                                )
                                continue
                            current_state["attempted_call_signatures"].add(signature)
                            resolved_calls.append((call, selected, call_input, call_context, signature))

                        if len(resolved_calls) == 0:
                            add_step(
                                {
                                    "status": "manager_no_valid_calls",
                                    "message": "Manager requested calls, but none matched existing agents.",
                                }
                            )
                        else:
                            current_state["agent_calls_count"] += len(resolved_calls)
                            add_step(
                                {
                                    "status": "manager_dispatch",
                                    "calls_dispatched": len(resolved_calls),
                                    "agent_calls_used": current_state["agent_calls_count"],
                                    "agent_calls_remaining": max(
                                        current_state["max_agent_calls"] - current_state["agent_calls_count"],
                                        0,
                                    ),
                                }
                            )

                            async def execute_call(
                                call: Dict[str, Any],
                                selected: Dict[str, Any],
                                call_input: str,
                                call_context: Dict[str, Any],
                                call_signature: str,
                            ) -> Dict[str, Any]:
                                selected_id = _safe_int(selected.get("id"))
                                add_step(
                                    {
                                        "status": "agent_call_started",
                                        "agent": selected.get("name"),
                                        "agent_id": selected_id,
                                        "agent_type": selected.get("agent_type"),
                                        "input": call_input,
                                        "query_name": call_context.get("query_name"),
                                        "params": call_context.get("params"),
                                    }
                                )

                                try:
                                    execution_config = build_agent_execution_config(selected)
                                    result = await AgentExecutor.execute(
                                        selected.get("agent_type") or "custom",
                                        execution_config,
                                        call_input,
                                        current_state["llm"],
                                        on_step=add_step,
                                        agent_name=selected.get("name") or selected.get("agent_type") or "agent",
                                        invocation=call_context if len(call_context) > 0 else None,
                                    )
                                    add_step(
                                        {
                                            "status": "agent_call_completed",
                                            "agent": selected.get("name"),
                                            "agent_id": selected_id,
                                            "result": result,
                                        }
                                    )
                                    if result.get("sql"):
                                        add_step({"status": "sql_generated", "sql": result["sql"]})
                                    answer_text = str(result.get("answer") or "").strip()
                                    terminal_reason = _extract_terminal_failure_reason(answer_text)
                                    if answer_text and not terminal_reason:
                                        current_state["successful_call_signatures"].add(call_signature)
                                    marked_unavailable = False

                                    if terminal_reason and selected_id is not None:
                                        per_agent = current_state["terminal_failure_counts"].setdefault(selected_id, {})
                                        reason_count = int(per_agent.get(terminal_reason, 0)) + 1
                                        per_agent[terminal_reason] = reason_count
                                        add_step(
                                            {
                                                "status": "agent_terminal_failure",
                                                "agent": selected.get("name"),
                                                "agent_id": selected_id,
                                                "reason": terminal_reason,
                                                "count": reason_count,
                                                "message": answer_text[:500],
                                            }
                                        )
                                        if reason_count >= _terminal_failure_threshold(terminal_reason):
                                            current_state["unavailable_agents"][selected_id] = (
                                                f"Repeated terminal failure ({terminal_reason}): {answer_text[:500]}"
                                            )
                                            marked_unavailable = True
                                            add_step(
                                                {
                                                    "status": "agent_marked_unavailable",
                                                    "agent": selected.get("name"),
                                                    "agent_id": selected_id,
                                                    "reason": current_state["unavailable_agents"][selected_id],
                                                }
                                            )

                                    if _should_mark_agent_unavailable(answer_text) and selected_id is not None:
                                        current_state["unavailable_agents"][selected_id] = answer_text[:500]
                                        marked_unavailable = True
                                        add_step(
                                            {
                                                "status": "agent_marked_unavailable",
                                                "agent": selected.get("name"),
                                                "agent_id": selected_id,
                                                "reason": answer_text,
                                            }
                                        )
                                    informative = bool(
                                        answer_text
                                        and not terminal_reason
                                        and not _is_generic_orchestration_text(answer_text)
                                    )
                                    return {
                                        "trace": f"\nAgent {selected.get('name')} responded: {result.get('answer')}\n",
                                        "informative": informative,
                                        "terminal_reason": terminal_reason,
                                        "marked_unavailable": marked_unavailable,
                                        "agent": selected.get("name"),
                                        "detail": answer_text[:500],
                                    }
                                except Exception as exc:
                                    error_text = str(exc)
                                    terminal_reason = _extract_terminal_failure_reason(error_text)
                                    marked_unavailable = False
                                    add_step(
                                        {
                                            "status": "agent_call_failed",
                                            "agent": selected.get("name"),
                                            "agent_id": selected_id,
                                            "error": error_text,
                                        }
                                    )
                                    if terminal_reason and selected_id is not None:
                                        per_agent = current_state["terminal_failure_counts"].setdefault(selected_id, {})
                                        reason_count = int(per_agent.get(terminal_reason, 0)) + 1
                                        per_agent[terminal_reason] = reason_count
                                        add_step(
                                            {
                                                "status": "agent_terminal_failure",
                                                "agent": selected.get("name"),
                                                "agent_id": selected_id,
                                                "reason": terminal_reason,
                                                "count": reason_count,
                                                "message": error_text[:500],
                                            }
                                        )
                                        if reason_count >= _terminal_failure_threshold(terminal_reason):
                                            current_state["unavailable_agents"][selected_id] = (
                                                f"Repeated terminal failure ({terminal_reason}): {error_text[:500]}"
                                            )
                                            marked_unavailable = True
                                            add_step(
                                                {
                                                    "status": "agent_marked_unavailable",
                                                    "agent": selected.get("name"),
                                                    "agent_id": selected_id,
                                                    "reason": current_state["unavailable_agents"][selected_id],
                                                }
                                            )
                                    if _should_mark_agent_unavailable(error_text) and selected_id is not None:
                                        current_state["unavailable_agents"][selected_id] = error_text[:500]
                                        marked_unavailable = True
                                        add_step(
                                            {
                                                "status": "agent_marked_unavailable",
                                                "agent": selected.get("name"),
                                                "agent_id": selected_id,
                                                "reason": error_text,
                                            }
                                        )
                                    return {
                                        "trace": f"\nAgent {selected.get('name')} failed: {exc}\n",
                                        "informative": False,
                                        "terminal_reason": terminal_reason,
                                        "marked_unavailable": marked_unavailable,
                                        "agent": selected.get("name"),
                                        "detail": error_text[:500],
                                    }

                            outcomes = await asyncio.gather(
                                *(
                                    execute_call(call, selected, call_input, call_context, signature)
                                    for call, selected, call_input, call_context, signature in resolved_calls
                                )
                            )
                            current_state["conversation_history"] += "".join(
                                str(item.get("trace") or "") for item in outcomes
                            )
                            current_state["flight_plan"] = _apply_plan_progress_from_outcomes(
                                current_state.get("flight_plan") or [],
                                outcomes,
                            )

                            try:
                                reflection_payload = [
                                    {
                                        "agent": item.get("agent"),
                                        "informative": bool(item.get("informative")),
                                        "terminal_reason": item.get("terminal_reason"),
                                        "marked_unavailable": bool(item.get("marked_unavailable")),
                                        "detail": _clip_text(str(item.get("detail") or ""), 400),
                                    }
                                    for item in outcomes
                                ]
                                reflection_prompt = f"""You are the reflection module of a multi-agent manager.
Evaluate the latest agent outcomes and decide whether the current plan needs adjustment.

Task objective:
{_clip_text(str(current_state.get("input") or ""), 500)}

Current plan:
{_clip_text(str(current_state.get("current_plan") or ""), 2000)}

Current flight plan (DAG):
{_plan_graph_to_text(current_state.get("flight_plan") or [])}

Latest outcomes:
{json.dumps(reflection_payload, ensure_ascii=False, indent=2)}

Return ONLY JSON:
{{
  "step_objective_met": true,
  "reflection": "short analysis",
  "plan_adjustment_needed": false,
  "updated_plan": "optional revised plan text",
  "updated_dag_plan": [{{"id":"T1","task":"...","depends_on":[],"status":"pending","success_criteria":"..."}}],
  "next_focus": "what to do next"
}}"""
                                reflection_response = await current_state["llm"].ainvoke(
                                    [HumanMessage(content=reflection_prompt)]
                                )
                                reflection_content = _coerce_message_content(reflection_response.content)
                                reflection = _safe_json_loads(reflection_content)
                                if reflection:
                                    updated_plan = _clip_text(str(reflection.get("updated_plan") or ""), 2600)
                                    if updated_plan:
                                        current_state["current_plan"] = updated_plan
                                        current_state["plan_history"].append(updated_plan)
                                    updated_dag = _normalize_plan_graph(reflection.get("updated_dag_plan"))
                                    if updated_dag:
                                        current_state["flight_plan"] = updated_dag
                                        current_state["plan_created"] = True
                                    reflection_text = _clip_text(str(reflection.get("reflection") or ""), 1200)
                                    next_focus = _clip_text(str(reflection.get("next_focus") or ""), 800)
                                    add_step(
                                        {
                                            "status": "manager_reflection",
                                            "step_objective_met": bool(reflection.get("step_objective_met")),
                                            "plan_adjustment_needed": bool(reflection.get("plan_adjustment_needed")),
                                            "reflection": reflection_text,
                                            "next_focus": next_focus,
                                        }
                                    )
                                    current_state["history_summary"] = _update_history_summary(
                                        current_state.get("history_summary") or "",
                                        (
                                            "Observation: "
                                            f"{reflection_text or 'No reflection text provided.'} "
                                            f"Next focus: {next_focus or 'Continue with current plan.'}"
                                        ),
                                    )
                            except Exception as reflection_exc:
                                add_step(
                                    {
                                        "status": "manager_reflection_failed",
                                        "error": str(reflection_exc),
                                    }
                                )

                            informative_outcomes = [item for item in outcomes if bool(item.get("informative"))]
                            if len(informative_outcomes) == 0:
                                terminal_outcomes = [
                                    item for item in outcomes if item.get("terminal_reason")
                                ]
                                if len(terminal_outcomes) > 0 and all(bool(item.get("marked_unavailable")) for item in terminal_outcomes):
                                    current_state["final_answer"] = _summarize_terminal_failures_for_user(
                                        [
                                            {
                                                "agent": item.get("agent"),
                                                "reason": item.get("terminal_reason"),
                                                "detail": item.get("detail"),
                                            }
                                            for item in terminal_outcomes
                                        ]
                                    )
                                    add_step(
                                        {
                                            "status": "manager_final",
                                            "answer": current_state["final_answer"],
                                            "manager_summary": "Stopped after repeated terminal agent failures.",
                                        }
                                    )
                                    current_state["done"] = True
                else:
                    current_state["conversation_history"] += f"\nManager thought: {rationale}\n"
                    current_state["history_summary"] = _update_history_summary(
                        current_state.get("history_summary") or "",
                        f"Step {current_state['current_step'] + 1}: no dispatch. rationale={_clip_text(rationale, 500)}",
                    )

            except Exception as exc:
                add_step({"status": "error", "message": f"Manager error: {exc}"})
                current_state["final_answer"] = f"Orchestration error: {exc}"
                current_state["done"] = True

            current_state["current_step"] += 1
            if current_state["current_step"] >= current_state["max_steps"] and not current_state["done"]:
                current_state["final_answer"] = "Reached maximum steps without a final answer."
                add_step(
                    {
                        "status": "manager_final",
                        "answer": current_state["final_answer"],
                        "manager_summary": "Max steps reached.",
                    }
                )
                current_state["done"] = True

            return current_state

        workflow = StateGraph(ManagerState)
        workflow.add_node("iterate", iterate)
        workflow.set_entry_point("iterate")
        workflow.add_conditional_edges("iterate", lambda s: END if s["done"] else "iterate")
        graph = workflow.compile()

        final_state = await graph.ainvoke(state)
        steps = final_state.get("steps")
        if not isinstance(steps, list):
            steps = []

        final_answer = str(final_state.get("final_answer") or "").strip()
        if not final_answer:
            final_answer = _extract_manager_answer_from_steps(steps)

        if not final_answer:
            try:
                synthesized = await _synthesize_manager_fallback_answer(
                    llm=final_state.get("llm"),
                    user_input=user_input,
                    conversation_history=str(final_state.get("conversation_history") or ""),
                    scratchpad=final_state.get("scratchpad") if isinstance(final_state.get("scratchpad"), dict) else {},
                    steps=steps,
                )
                if synthesized:
                    final_answer = synthesized
                    steps.append(
                        {
                            "status": "manager_finalized_fallback",
                            "answer": final_answer,
                            "manager_summary": "Fallback synthesis generated because manager did not provide a final answer.",
                        }
                    )
            except Exception as exc:
                steps.append(
                    {
                        "status": "manager_finalize_failed",
                        "error": str(exc),
                        "message": "Failed to synthesize fallback final answer.",
                    }
                )

        if not final_answer:
            final_answer = "Manager completed orchestration but did not produce a final answer."
            steps.append(
                {
                    "status": "manager_finalized_fallback",
                    "answer": final_answer,
                    "manager_summary": "Default fallback answer used.",
                }
            )

        return {
            "answer": final_answer,
            "steps": steps,
        }
